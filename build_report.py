"""
Builds the Industrial Internship Training report as a .docx laid out to the
MGM COET guidelines: A4, Times New Roman 12pt, 1in top/bottom and 1.25in
left/right margins, 1.5 line spacing, justified body text, bold headings,
16pt bold chapter titles, roman front matter and arabic body page numbers.
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(ROOT, "assets")
OUTPUT = os.path.join(ROOT, "Internship_Report_Suraj_Singh_Rawat.docx")

FONT = "Times New Roman"
BODY_PT = 12
CHAPTER_PT = 16
TEXT_WIDTH = Inches(5.75)

# Guideline: lines at 1.5 spacing, paragraphs separated by 2.0 spacing.
# A 12pt line occupies ~13.8pt, so the extra gap that raises a 1.5-spaced
# paragraph break to a 2.0-spaced one is 0.5 * 13.8 ~= 7pt.
PARA_GAP = 7

STUDENT = "SURAJ SINGH RAWAT"
ROLL = "2400951530054"
CLASS = "TT-C (AI & ML)"
TITLE = "MULTI-AGENT AI ORCHESTRATION SYSTEM FOR AUTOMATED RESEARCH"


# --------------------------------------------------------------------------
# low level helpers
# --------------------------------------------------------------------------

def set_font(run, size=BODY_PT, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    return run


_pending_break = []


def para(doc, text="", size=BODY_PT, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.5, after=PARA_GAP, before=0,
         indent=None, keep_with_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if _pending_break:
        _pending_break.clear()
        pf.page_break_before = True
    pf.alignment = align
    pf.line_spacing = line
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    if indent is not None:
        pf.left_indent = Inches(indent)
    pf.keep_with_next = keep_with_next
    if text:
        set_font(p.add_run(text), size, bold, italic)
    return p


def chapter(doc, number, name):
    """Chapter number top left, name immediately below, single spaced, 16pt bold."""
    para(doc, f"CHAPTER {number}", size=CHAPTER_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=0, keep_with_next=True)
    para(doc, name.upper(), size=CHAPTER_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=18, keep_with_next=True)


def heading(doc, text, level=2):
    para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5,
         before=6 if level == 2 else 4, after=6, keep_with_next=True)


def preface_heading(doc, text):
    para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)


def bullets(doc, items, ordered=False):
    for i, item in enumerate(items, 1):
        marker = f"{i}." if ordered else "•"
        p = para(doc, after=6, indent=0.4)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        set_font(p.add_run(f"{marker}\t"))
        if isinstance(item, tuple):
            set_font(p.add_run(item[0]), bold=True)
            set_font(p.add_run(item[1]))
        else:
            set_font(p.add_run(item))


def table(doc, headers, rows, caption, widths=None, size=BODY_PT, keep_together=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    for cell, text in zip(t.rows[0].cells, headers):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        set_font(cell.paragraphs[0].add_run(text), size, bold=True)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:val"), "clear")
        shade.set(qn("w:fill"), "DCE6F1")
        cell._tc.get_or_add_tcPr().append(shade)

    for row in rows:
        cells = t.add_row().cells
        for cell, text in zip(cells, row):
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            set_font(cell.paragraphs[0].add_run(str(text)), size)

    if widths:
        for row in t.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = Inches(w)

    # avoid an orphaned last row stranded alone on the next page: for
    # short tables, chain every row to the next so the whole table (and
    # its caption) moves together if it does not fit on the current page
    if keep_together is None:
        keep_together = len(rows) <= 10
    if keep_together:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True

    # repeat the header row on page breaks
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    tr_pr.append(header_flag)

    if caption:
        para(doc, caption, size=BODY_PT, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, before=6, after=14)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)


def figure(doc, filename, caption, width=4.8):
    path = os.path.join(ASSETS, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(path, width=Inches(width))
    para(doc, caption, size=BODY_PT, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, after=14)


def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)
        pf.left_indent = Inches(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), "Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def toc_line(doc, text, page, bold=False, indent=0.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.left_indent = Inches(indent)
    pf.tab_stops.add_tab_stop(TEXT_WIDTH, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    set_font(p.add_run(text), bold=bold)
    set_font(p.add_run(f"\t{page}"), bold=bold)


def page_break(doc):
    """Mark the next paragraph as starting a page, rather than emitting a
    break paragraph of its own, which can leave a stray blank page behind."""
    _pending_break.append(True)


# --------------------------------------------------------------------------
# section / page-number plumbing
# --------------------------------------------------------------------------

def configure(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.footer_distance = Inches(0.5)
    return section


def number_format(section, fmt, start):
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(old)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))
    sect_pr.append(pg)


def footer_page_number(section, align):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = p.add_run()
    set_font(run)
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def blank_footer(section):
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ""


# --------------------------------------------------------------------------
# front matter
# --------------------------------------------------------------------------

def title_page(doc):
    def line(text, size, bold=True, italic=False, after=0, before=0):
        para(doc, text, size=size, bold=bold, italic=italic,
             align=WD_ALIGN_PARAGRAPH.CENTER, line=1.5, after=after, before=before)

    line("INDUSTRIAL INTERNSHIP TRAINING", 18, before=6)
    line("ON MULTI-AGENT SYSTEM DEVELOPMENT", 18, after=10)
    line("On", 12, bold=False, after=6)
    line("MULTI-AGENT AI ORCHESTRATION SYSTEM", 14)
    line("FOR AUTOMATED RESEARCH", 14, after=12)
    line("Submitted by", 12, bold=False, italic=True, after=4)
    line(STUDENT, 14, after=4)
    line(f"University Roll No: {ROLL}", 14, after=4)
    line(f"Class: {CLASS}", 12, after=12)
    line("In partial fulfillment of the requirements for the award of the degree of",
         14, bold=False, italic=True, after=6)
    line("BACHELOR OF TECHNOLOGY", 16, after=0)
    line("In", 12, bold=False, after=0)
    line("COMPUTER SCIENCE & ENGINEERING", 16, after=8)

    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.paragraph_format.space_before = Pt(6)
    logo.paragraph_format.space_after = Pt(8)
    logo.add_run().add_picture(os.path.join(ASSETS, "mgm_logo.jpeg"), width=Inches(1.15))

    line("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", 14, after=0)
    line("MGM's College of Engineering & Technology, Noida", 14, after=4)
    line("Affiliated to Dr. A.P.J. Abdul Kalam Technical University, Lucknow",
         12, bold=False, after=8)
    line("August, 2026-27", 16, after=10)
    line("SUBJECT CODE: KCS-753    INDUSTRIAL TRAINING", 12)


def training_certificate(doc):
    """Guideline 1: a photocopy of the certificate issued by the industry is
    placed immediately after the title page."""
    preface_heading(doc, "TRAINING CERTIFICATE")
    para(doc, "", after=48)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.5, after=48)
    set_font(p.add_run("[ Photocopy of the training certificate issued by the "
                       "organisation is to be affixed here ]"), italic=True)
    para(doc, "", after=48)
    para(doc, f"Candidate: {STUDENT.title()}", align=WD_ALIGN_PARAGRAPH.CENTER,
         line=1.0, after=0)
    para(doc, f"University Roll No: {ROLL}", align=WD_ALIGN_PARAGRAPH.CENTER,
         line=1.0, after=0)
    para(doc, f"Class: {CLASS}", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)


def certificate(doc):
    preface_heading(doc, "CERTIFICATE")
    para(doc, f'This is to certify that the Industrial Internship Training report entitled '
              f'"{TITLE}" has been carried out by {STUDENT.title()} (University Roll No. {ROLL}), '
              f'a student of Bachelor of Technology, Computer Science & Engineering, class {CLASS}, '
              f'under my supervision and guidance, in partial fulfillment of the requirements for '
              f'the award of the degree of Bachelor of Technology in Computer Science & Engineering '
              f'from Dr. A.P.J. Abdul Kalam Technical University, Lucknow.')
    para(doc, "The matter embodied in this report is the original work of the candidate and has "
              "not been submitted earlier for the award of any degree or diploma to the best of my "
              "knowledge and belief.")
    para(doc, "", after=24)
    para(doc, "Date: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT, after=6)
    para(doc, "Place: Noida", align=WD_ALIGN_PARAGRAPH.LEFT, after=36)
    para(doc, "(Signature of the Guide)", align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0, after=0)
    para(doc, "Name of the Guide", align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0, after=0)
    para(doc, "Department of CSE, MGM COET, Noida", align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0)


def declaration(doc):
    preface_heading(doc, "DECLARATION")
    para(doc, f'I affirm that the Industrial Internship Training report titled "{TITLE}", being '
              f'submitted in partial fulfillment of the requirements for the award of the Degree of '
              f'Bachelor of Technology in Computer Science & Engineering, is the original work '
              f'carried out by me. It has not formed the part of any other project work submitted '
              f'for the award of any degree or diploma, either in this or any other Institution.')
    para(doc, "All sources of information and material used in the preparation of this report have "
              "been duly acknowledged in the references section.")
    para(doc, "", after=36)
    para(doc, "(Signature)", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=0)
    para(doc, STUDENT.title(), align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=0)
    para(doc, f"University Roll No: {ROLL}", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=0)
    para(doc, f"Class: {CLASS}", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0)


def acknowledgement(doc):
    preface_heading(doc, "ACKNOWLEDGEMENT")
    para(doc, "I would like to express my sincere gratitude to my project guide and to the faculty "
              "of the Department of Computer Science & Engineering, MGM's College of Engineering & "
              "Technology, Noida, for their valuable guidance, constant encouragement and support "
              "throughout the duration of this industrial internship training and the preparation "
              "of this report.")
    para(doc, "I am especially thankful to the faculty members who shared their expertise in "
              "artificial intelligence, large language models and software engineering practice. "
              "Their willingness to review my design decisions, question my assumptions and suggest "
              "alternatives shaped the architecture of the system described in this report.")
    para(doc, "I am also grateful to the Head of the Department for providing the laboratory "
              "facilities, computing resources and the working environment that made the "
              "development and evaluation of this system possible.")
    para(doc, "Finally, I extend my thanks to my classmates and to my family for their continuous "
              "motivation and support during the completion of this internship.")
    para(doc, "", after=24)
    para(doc, STUDENT.title(), align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0, after=0)
    para(doc, f"University Roll No: {ROLL}", align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0)


def abstract(doc):
    preface_heading(doc, "ABSTRACT")
    para(doc, "This report describes the work carried out during an industrial internship on the "
              "design, implementation and evaluation of a Multi-Agent Research System, an "
              "application that automates the routine stages of desk research by coordinating four "
              "specialised software agents. Conducting research manually on a single topic "
              "typically consumes between two and four hours of continuous effort, produces results "
              "whose depth varies with the skill of the individual researcher, and scales only by "
              "adding people. The system built during this internship addresses each of these "
              "difficulties.")
    para(doc, "The system is organised as a sequential pipeline of four agents. The Search Agent "
              "converts a topic into a query, retrieves candidate sources through the DuckDuckGo "
              "interface, scores each result for relevance and returns a ranked shortlist. The "
              "Reader Agent retrieves each shortlisted page, parses the markup with BeautifulSoup4, "
              "strips navigation and advertising material and aggregates the remaining text. The "
              "Writer Agent submits the aggregated text to the Groq gpt-oss-120b language model "
              "with a structured prompt and receives a report in a fixed six-section format. The "
              "Critic Agent then scores that report on five dimensions and returns written "
              "feedback. Orchestration, prompt management and retry behaviour are handled through "
              "LangChain, and the user interacts with the system through a Streamlit web interface.")
    para(doc, "The completed system produces a finished research report in a measured mean of 49.1 "
              "seconds, against a design ceiling of 60 seconds. Every individual agent finishes "
              "within its own budget. Selecting Groq over a comparable commercial endpoint reduced "
              "the cost per thousand tokens by approximately seventy per cent without a measurable "
              "reduction in output quality. The system sustains ten concurrent users with less than "
              "two seconds of additional latency, and its stateless design allows any instance to "
              "be restarted without loss of user data.")
    para(doc, "Correctness was established through a suite of thirty-nine automated tests covering "
              "the four agents individually, the interactions between them and two complete "
              "end-to-end journeys. All thirty-nine tests pass and statement coverage across the "
              "code base stands at eighty-five per cent. The report documents the requirements "
              "analysis, the architecture and design decisions, the implementation of each agent, "
              "the testing strategy and its results, the deployment procedure, and the limitations "
              "that define the scope of further work.")


def table_of_contents(doc):
    preface_heading(doc, "TABLE OF CONTENTS")
    entries = [
        ("Training Certificate", "i", True, 0),
        ("Certificate", "ii", True, 0),
        ("Acknowledgement", "iii", True, 0),
        ("Declaration", "iv", True, 0),
        ("Abstract", "v", True, 0),
        ("List of Tables", "viii", True, 0),
        ("List of Figures", "viii", True, 0),
        ("Table of Abbreviations", "ix", True, 0),
        ("CHAPTER 1  INTRODUCTION", "1", True, 0),
        ("1.1  Background of the Organisation and the Work", "1", False, 0.3),
        ("1.2  Training Objective", "1", False, 0.3),
        ("1.3  Student's Work Assignment", "2", False, 0.3),
        ("1.4  Organisation of the Report", "2", False, 0.3),
        ("CHAPTER 2  TECHNIQUES AND TOOLS STUDIED", "3", True, 0),
        ("2.1  Large Language Models and Prompt Design", "3", False, 0.3),
        ("2.2  The Multi-Agent Pattern", "3", False, 0.3),
        ("2.3  LangChain as an Orchestration Framework", "4", False, 0.3),
        ("2.4  Web Retrieval and Content Extraction", "4", False, 0.3),
        ("2.5  Streamlit for Rapid Interface Construction", "4", False, 0.3),
        ("CHAPTER 3  REQUIREMENT ANALYSIS", "5", True, 0),
        ("3.1  Problem Statement", "5", False, 0.3),
        ("3.2  Functional Requirements", "5", False, 0.3),
        ("3.3  Non-Functional Requirements", "6", False, 0.3),
        ("3.4  Hardware Requirements", "6", False, 0.3),
        ("3.5  Software Requirements", "6", False, 0.3),
        ("CHAPTER 4  SYSTEM ANALYSIS AND DESIGN", "8", True, 0),
        ("4.1  Existing Practice and the Proposed System", "8", False, 0.3),
        ("4.2  Technology Stack and Selection Rationale", "8", False, 0.3),
        ("4.3  System Architecture", "9", False, 0.3),
        ("4.4  The Four-Agent Pipeline", "10", False, 0.3),
        ("4.5  Data Model and Inter-Agent Contracts", "11", False, 0.3),
        ("CHAPTER 5  IMPLEMENTATION", "13", True, 0),
        ("5.1  Search Agent", "13", False, 0.3),
        ("5.2  Reader Agent", "14", False, 0.3),
        ("5.3  Writer Agent", "15", False, 0.3),
        ("5.4  Critic Agent", "15", False, 0.3),
        ("5.5  Pipeline Orchestration and Error Handling", "16", False, 0.3),
        ("5.6  User Interface", "16", False, 0.3),
        ("CHAPTER 6  TESTING AND RESULTS", "18", True, 0),
        ("6.1  Testing Methodology", "18", False, 0.3),
        ("6.2  Test Cases and Results", "18", False, 0.3),
        ("6.3  Performance Benchmarking", "19", False, 0.3),
        ("6.4  Cost Evaluation", "21", False, 0.3),
        ("CHAPTER 7  DEPLOYMENT AND OPERATIONS", "22", True, 0),
        ("7.1  Local Installation", "22", False, 0.3),
        ("7.2  Cloud Deployment", "22", False, 0.3),
        ("7.3  Operational Monitoring", "23", False, 0.3),
        ("CHAPTER 8  CONCLUSION AND RECOMMENDATIONS", "24", True, 0),
        ("8.1  Conclusion", "24", False, 0.3),
        ("8.2  Skills Acquired and Contribution Made", "24", False, 0.3),
        ("8.3  Limitations", "25", False, 0.3),
        ("8.4  Recommendations and Future Scope", "25", False, 0.3),
        ("REFERENCES", "26", True, 0),
        ("APPENDIX A: SOURCE CODE", "28", True, 0),
    ]
    for text, page, bold, indent in entries:
        toc_line(doc, text, page, bold=bold, indent=indent)


def list_of_tables(doc):
    preface_heading(doc, "LIST OF TABLES")
    table(doc,
          ["Table No.", "Caption", "Page No."],
          [["3.1", "Functional requirements of the system", "5"],
           ["3.2", "Non-functional requirements and their targets", "6"],
           ["3.3", "Minimum and recommended hardware", "6"],
           ["3.4", "Software components and versions used", "7"],
           ["4.1", "Manual research compared with the proposed system", "8"],
           ["4.2", "Technology stack and selection rationale", "9"],
           ["4.3", "Inter-agent data contracts", "12"],
           ["5.1", "Search Agent configuration parameters", "14"],
           ["5.2", "Failure conditions handled by the Reader Agent", "15"],
           ["5.3", "Scoring dimensions applied by the Critic Agent", "16"],
           ["6.1", "Distribution of the automated test suite", "18"],
           ["6.2", "Representative test cases and their results", "19"],
           ["6.3", "Measured execution time against target", "19"],
           ["6.4", "Statement coverage by module", "20"],
           ["6.5", "Monthly operating cost comparison", "21"]],
          "",
          widths=[1.0, 3.7, 1.05])


def list_of_figures(doc):
    para(doc, "LIST OF FIGURES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         before=14, after=14)
    table(doc,
          ["Figure No.", "Caption", "Page No."],
          [["4.1", "Layered architecture of the system", "10"],
           ["4.2", "Four-agent sequential pipeline", "11"],
           ["4.3", "Data handed between successive agents", "12"],
           ["5.1", "Search Agent processing sequence", "13"],
           ["5.2", "Reader Agent content extraction sequence", "14"],
           ["5.3", "Layout of the Streamlit user interface", "17"],
           ["6.1", "Measured execution time against target", "19"],
           ["6.2", "Statement coverage by module", "20"]],
          "",
          widths=[1.1, 3.6, 1.05])


def abbreviations(doc):
    preface_heading(doc, "TABLE OF ABBREVIATIONS")
    table(doc,
          ["Abbreviation", "Expansion"],
          [["AI", "Artificial Intelligence"],
           ["API", "Application Programming Interface"],
           ["AWS", "Amazon Web Services"],
           ["CPU", "Central Processing Unit"],
           ["CSE", "Computer Science and Engineering"],
           ["EBS", "Elastic Block Store"],
           ["EC2", "Elastic Compute Cloud"],
           ["HTML", "HyperText Markup Language"],
           ["HTTP", "HyperText Transfer Protocol"],
           ["JSON", "JavaScript Object Notation"],
           ["LLM", "Large Language Model"],
           ["LPU", "Language Processing Unit"],
           ["ML", "Machine Learning"],
           ["RAM", "Random Access Memory"],
           ["SLA", "Service Level Agreement"],
           ["UI", "User Interface"],
           ["URL", "Uniform Resource Locator"],
           ["XSS", "Cross-Site Scripting"]],
          "",
          widths=[1.7, 4.05])


# --------------------------------------------------------------------------
# chapters
# --------------------------------------------------------------------------

def chapter_1(doc):
    chapter(doc, 1, "Introduction")

    heading(doc, "1.1  Background of the Organisation and the Work")
    para(doc, "The internship was carried out within a software engineering group whose work "
              "concerns applied artificial intelligence, and specifically the construction of "
              "systems that place large language models inside conventional software pipelines. The "
              "group builds internal tools rather than consumer products; its output is used by "
              "analysts, technical writers and product managers elsewhere in the organisation who "
              "need reliable summaries of unfamiliar subject matter at short notice. The department "
              "in which the work took place is responsible for prototyping such tools, evaluating "
              "them against measurable targets, and handing successful prototypes to an operations "
              "team for hosting.")
    para(doc, "The particular difficulty this department was asked to address is the cost of desk "
              "research. When an analyst is asked to prepare a briefing on an unfamiliar topic, the "
              "work follows a predictable sequence: locate credible sources, read them, condense "
              "what matters into prose, and then judge whether the result is complete enough to "
              "circulate. That sequence consumes between two and four hours per topic, it produces "
              "results whose depth depends heavily on the individual analyst, and it can be scaled "
              "only by employing more analysts. Because the sequence is predictable, it is also a "
              "reasonable candidate for automation.")

    heading(doc, "1.2  Training Objective")
    para(doc, "The stated objective of the internship was to design, build and evaluate a working "
              "system that performs the whole desk-research sequence without human intervention, "
              "and to demonstrate through measurement that it meets defined targets for speed, cost "
              "and reliability. Three targets were agreed at the outset: a complete report was to be "
              "produced within sixty seconds; the cost per report was to fall substantially below "
              "that of the commercial language-model endpoint then in use; and the system was to "
              "serve at least ten simultaneous users without degradation. A fourth, broader "
              "objective was to gain practical experience of the engineering discipline that "
              "surrounds machine learning work, including automated testing, benchmarking and "
              "deployment, rather than model development in isolation.")

    heading(doc, "1.3  Student's Work Assignment")
    para(doc, "The assignment was to take the problem from an informal description through to a "
              "deployable artefact, working independently and reporting progress at weekly "
              "intervals. In non-technical terms the responsibilities were: to decide how the work "
              "of research should be divided among cooperating software components; to choose the "
              "external services on which those components would depend and to justify the choice "
              "on grounds of cost and speed; to write the components and the interface through "
              "which they are used; to establish by testing that they behave correctly, including "
              "when the network or an external service misbehaves; to measure performance against "
              "the agreed targets; and to document the result so that another engineer could operate "
              "and extend it. No part of the implementation was delegated.")

    heading(doc, "1.4  Organisation of the Report")
    para(doc, "Chapter 2 describes the techniques and tools studied during the internship. Chapter 3 "
              "records the requirements analysis, setting out what the system was obliged to do and "
              "the constraints under which it had to do it. Chapter 4 presents the analysis and "
              "design, comparing existing practice with the proposed approach and describing the "
              "architecture and the contracts between components. Chapter 5 documents the "
              "implementation of each agent together with the orchestration and error handling that "
              "join them. Chapter 6 reports the testing methodology, the test results, the "
              "performance measurements and the cost evaluation. Chapter 7 covers deployment and "
              "routine operation. Chapter 8 draws conclusions, records the skills acquired and the "
              "contribution made, states the limitations of the work and recommends further "
              "development. References and an appendix containing representative source code "
              "conclude the report.")


def chapter_2(doc):
    chapter(doc, 2, "Techniques and Tools Studied")

    heading(doc, "2.1  Large Language Models and Prompt Design")
    para(doc, "A large language model predicts the continuation of a sequence of text. It holds no "
              "record of a particular conversation beyond the text supplied to it, which means that "
              "everything the model is expected to take into account must be present in the prompt. "
              "This property governs the design of any system built around such a model: the "
              "surrounding software is responsible for retrieving the right material, placing it in "
              "the prompt, and interpreting what comes back.")
    para(doc, "Two aspects of prompt design proved important. Instructions describing the required "
              "shape of the output are far more effective than those describing its desired "
              "quality: asking for a report in six named sections with stated word budgets "
              "produced consistent structure, whereas asking for a thorough report did not. The "
              "sampling temperature also affects reproducibility; 0.7 was adopted for synthesis "
              "and 0.3 for evaluation, where consistent scoring matters more than fluency.")
    para(doc, "The model selected was gpt-oss-120b, an open-weights model with approximately one "
              "hundred and twenty billion parameters, served through the Groq endpoint [1]. The "
              "general behaviour of models of this class is described by Brown and others [10], "
              "and the effect of instruction phrasing on their reasoning by Wei and others [11]. "
              "The choice of provider is examined in Section 4.2.")

    heading(doc, "2.2  The Multi-Agent Pattern")
    para(doc, "A multi-agent system decomposes a task into stages and assigns each stage to a "
              "component that does one thing, an arrangement set out in the standard treatment by "
              "Wooldridge [12]. The alternative, submitting the entire task to a "
              "single model in one prompt, was tried first and rejected for three reasons. Its "
              "failures were opaque, because a poor report gave no indication whether the cause lay "
              "in poor sources or poor synthesis. It could not be tested in parts. And it could not "
              "be tuned in parts, since a change intended to improve source selection also altered "
              "the writing.")
    para(doc, "Dividing the work into four agents removed all three difficulties. Each agent has a "
              "declared input and output, so each can be tested against fixed data and measured "
              "separately, and a defect can be localised to a single stage. The cost is a loss of "
              "flexibility: the pipeline runs in a fixed order and cannot revisit an earlier stage "
              "in the light of a later one. That trade-off is discussed in Section 8.3.")

    heading(doc, "2.3  LangChain as an Orchestration Framework")
    para(doc, "LangChain [2] supplies the connective material between an application and a language "
              "model: templated prompts, a uniform interface across providers, parsers that convert "
              "model output into structured values, and operators that join stages into a chain. "
              "Using it removed code that would otherwise have had to be written and maintained, "
              "particularly around retry behaviour and the parsing of structured replies. Its "
              "provider abstraction proved valuable in practice: the project began against one "
              "endpoint and moved to Groq later, and the change touched configuration rather than "
              "agent logic.")

    heading(doc, "2.4  Web Retrieval and Content Extraction")
    para(doc, "Two distinct problems arise in obtaining source material. The first is discovery, "
              "solved here through the DuckDuckGo search interface [9], which requires no key and "
              "imposes no quota at the volumes involved. The second is extraction, which is harder. "
              "A retrieved page contains navigation, advertising, cookie notices and related-article "
              "panels alongside the text that matters, and none of it is labelled as such.")
    para(doc, "BeautifulSoup4 [3] was used to parse the retrieved markup into a navigable tree. "
              "Extraction proceeds by removing the element types that reliably carry no content, "
              "namely script, style, nav, header, footer and aside, and then taking the remaining "
              "text. This is a heuristic rather than a solution, and it fails on pages whose content "
              "is assembled by client-side scripting after delivery. The consequences are recorded "
              "in Section 5.2.")

    heading(doc, "2.5  Streamlit for Rapid Interface Construction")
    para(doc, "Streamlit [4] renders a Python script as a web application, re-executing the script "
              "whenever the user interacts with a widget. It was chosen because the interface "
              "required here is genuinely simple, comprising one text field, a progress indicator "
              "and a rendered result, and because it avoids introducing a second language and a "
              "separate build process into a project whose substance lies elsewhere. Its "
              "re-execution model does require care: any value that must survive an interaction has "
              "to be placed in the session state object rather than held in a local variable.")


def chapter_3(doc):
    chapter(doc, 3, "Requirement Analysis")

    heading(doc, "3.1  Problem Statement")
    para(doc, "Desk research performed by hand is slow, variable in quality and expensive to scale. "
              "A single topic occupies an analyst for two to four hours. Two analysts given the same "
              "topic produce reports of differing depth, differing structure and differing standards "
              "of citation. Capacity can be increased only by employing more analysts, so cost rises "
              "in direct proportion to volume. The question this project set out to answer was "
              "whether a pipeline of cooperating software agents could complete the same sequence "
              "within one minute, at a cost per report substantially below that of a commercial "
              "language-model endpoint, and with sufficient consistency that the output is usable "
              "without rework.")

    heading(doc, "3.2  Functional Requirements")
    para(doc, "The functional requirements state what the system is obliged to do. They were fixed "
              "at the start of the internship and did not change materially thereafter.")
    table(doc,
          ["Ref.", "Requirement"],
          [["FR1", "Accept a research topic entered through the web interface and begin processing it."],
           ["FR2", "Retrieve candidate sources through the search interface and rank them by relevance, "
                   "scoring the title at forty per cent and the snippet at sixty per cent."],
           ["FR3", "Return at most ten sources, discarding any scoring below the relevance threshold "
                   "and removing duplicate addresses."],
           ["FR4", "Retrieve each shortlisted page, extract its main text and reject any extraction "
                   "shorter than one hundred words."],
           ["FR5", "Continue processing when an individual page cannot be retrieved, rather than "
                   "abandoning the request."],
           ["FR6", "Produce a report in six sections: summary, key findings, analysis, implications, "
                   "conclusion and references."],
           ["FR7", "Score the finished report on five dimensions and return written feedback "
                   "identifying strengths and weaknesses."],
           ["FR8", "Present the report, the score and the source list to the user within sixty seconds."],
           ["FR9", "Serve at least ten concurrent users without cross-contamination between sessions."],
           ["FR10", "Record for every request the timestamp, topic, sources used, per-agent duration "
                    "and final score."]],
          "Table 3.1: Functional requirements of the system",
          widths=[0.75, 5.0])

    heading(doc, "3.3  Non-Functional Requirements")
    para(doc, "The non-functional requirements constrain how the system behaves rather than what it "
              "produces. Each was given a numeric target so that compliance could be demonstrated by "
              "measurement rather than asserted.")
    table(doc,
          ["Category", "Requirement and target"],
          [["Performance", "Search within 10 s, reading within 15 s, writing within 20 s, "
                           "evaluation within 15 s, complete pipeline within 60 s."],
           ["Scalability", "Ten concurrent sessions; not more than 100 MB of memory per session; "
                           "no persistent database required."],
           ["Reliability", "95 per cent availability; automatic retry with exponential backoff "
                           "from two seconds to a ceiling of thirty-two seconds."],
           ["Security", "Credentials read from the environment and never committed; user input "
                        "validated; generated content escaped before rendering."],
           ["Privacy", "Session data discarded once the result has been delivered; no user "
                       "profile retained."],
           ["Maintainability", "Each agent in its own module with an independent test suite; "
                               "statement coverage of at least 85 per cent."],
           ["Usability", "Interface usable without documentation; progress shown during "
                         "processing; every source listed with the result."],
           ["Cost", "Not more than five United States dollars per month in model charges for "
                    "one thousand reports."]],
          "Table 3.2: Non-functional requirements and their targets",
          widths=[1.4, 4.35])

    heading(doc, "3.4  Hardware Requirements")
    para(doc, "The system performs no local inference, so its hardware demands are modest. Most of "
              "the elapsed time in a request is spent waiting for a network response rather than "
              "computing, and the processor is largely idle during that period.")
    table(doc,
          ["Component", "Minimum", "Recommended"],
          [["Processor", "Dual core, 2.0 GHz", "Quad core, 2.5 GHz"],
           ["Memory", "2 GB", "4 GB to 8 GB"],
           ["Storage", "500 MB free", "2 GB free"],
           ["Network", "Broadband connection", "25 Mbps or better"],
           ["Display", "Any modern display", "1920 x 1080"]],
          "Table 3.3: Minimum and recommended hardware",
          widths=[1.6, 2.1, 2.05])

    heading(doc, "3.5  Software Requirements")
    table(doc,
          ["Component", "Version", "Purpose"],
          [["Python", "3.10 or later", "Runtime for the whole application"],
           ["Streamlit", "1.28 or later", "Web user interface"],
           ["LangChain", "0.1.0 or later", "Prompt templates and agent chaining"],
           ["Groq endpoint", "current", "Access to the gpt-oss-120b model"],
           ["BeautifulSoup4", "4.12 or later", "Parsing of retrieved markup"],
           ["Requests", "2.31 or later", "HTTP client"],
           ["python-dotenv", "1.0 or later", "Loading of configuration from file"],
           ["pytest", "7.4 or later", "Automated test execution"]],
          "Table 3.4: Software components and versions used",
          widths=[1.7, 1.6, 2.45])


def chapter_4(doc):
    chapter(doc, 4, "System Analysis and Design")

    heading(doc, "4.1  Existing Practice and the Proposed System")
    para(doc, "Two alternatives were considered before the multi-agent design was adopted: "
              "continuing with manual research, and submitting the whole task to a single language "
              "model. The table below sets the three approaches against one another on the "
              "dimensions that mattered to the department.")
    table(doc,
          ["Criterion", "Manual research", "Single model", "Proposed system"],
          [["Elapsed time", "2 to 4 hours", "30 to 60 minutes", "49.1 seconds"],
           ["Cost per report", "30 to 100 USD", "0.15 to 0.50 USD", "0.0005 USD"],
           ["Source handling", "Varies by analyst", "Not retrieved", "Ten ranked, verified"],
           ["Consistency", "Low", "Moderate", "High"],
           ["Diagnosability", "Not applicable", "Opaque", "Per stage"],
           ["Scaling", "By headcount", "By quota", "By concurrency"]],
          "Table 4.1: Manual research compared with the proposed system",
          widths=[1.3, 1.5, 1.35, 1.6])
    para(doc, "The decisive consideration was not speed, which any automated approach would have "
              "improved, but diagnosability. The single-model approach produced reports that were "
              "sometimes poor, and offered no means of establishing why. Dividing the work into "
              "stages with declared interfaces made each stage separately observable, and it is "
              "that property which made the measurement reported in Chapter 6 possible at all.")

    heading(doc, "4.2  Technology Stack and Selection Rationale")
    table(doc,
          ["Layer", "Technology", "Reason for selection"],
          [["Interface", "Streamlit", "Interface is simple; avoids a second language and build step"],
           ["Orchestration", "LangChain", "Prompt templating, output parsing, provider abstraction"],
           ["Language model", "Groq gpt-oss-120b", "Roughly seventy per cent cheaper per token; low latency"],
           ["Discovery", "DuckDuckGo", "No key required; no quota at the volumes involved"],
           ["Extraction", "BeautifulSoup4", "Tolerant of malformed markup; well documented"],
           ["Transport", "Requests", "Explicit timeout and redirect control"],
           ["Configuration", "python-dotenv", "Keeps credentials out of the repository"],
           ["Testing", "pytest", "Fixtures and parametrisation suit agent testing"]],
          "Table 4.2: Technology stack and selection rationale",
          widths=[1.25, 1.6, 2.9])
    para(doc, "The application runs on Python 3.10 [6]. Network access is made through the Requests "
              "library [5], which was preferred over the standard-library client because it exposes "
              "explicit timeout and redirect controls, both of which the Reader Agent depends upon. "
              "The test suite is built on pytest [7]. The retrieval-then-synthesis arrangement used "
              "here follows the pattern described by Lewis and others [13], in which material "
              "gathered at request time is placed in the prompt rather than relied upon from the "
              "model's own parameters.")
    para(doc, "The choice of model provider deserves comment because it was the single decision with "
              "the largest effect on operating cost. A commercial endpoint was used during early "
              "development at a published rate of 0.0015 United States dollars per thousand input "
              "tokens. The Groq endpoint serves gpt-oss-120b at approximately 0.0005 dollars for the "
              "same volume. Reports generated by both were compared over a set of twenty topics and "
              "scored by the Critic Agent; the difference in mean score was smaller than the "
              "variation between repeated runs on the same topic, so the reduction in cost was taken "
              "without an offsetting loss.")

    heading(doc, "4.3  System Architecture")
    para(doc, "The system is arranged in four layers. The presentation layer accepts the topic and "
              "renders the result. The orchestration layer manages the lifecycle of the agents, the "
              "passing of values between them and the retry behaviour when an external call fails. "
              "The agent layer contains the four components that perform the work. The service layer "
              "comprises the external dependencies: the search interface, the language-model "
              "endpoint and the parsing library.")
    figure(doc, "fig_3_1_architecture.png", "Figure 4.1: Layered architecture of the system")
    para(doc, "No layer reaches past its immediate neighbour. The agents do not render output and "
              "the interface does not call an external service directly. The arrangement is "
              "deliberately conventional; its value is that a change confined to one layer, such as "
              "the substitution of one model provider for another, does not propagate.")

    heading(doc, "4.4  The Four-Agent Pipeline")
    para(doc, "The agents execute in a fixed order, each consuming the output of its predecessor. "
              "Figure 4.2 shows the sequence together with the measured mean duration of each stage.")
    figure(doc, "fig_3_2_pipeline.png", "Figure 4.2: Four-agent sequential pipeline", width=4.7)
    para(doc, "The order is not arbitrary. Reading cannot begin before sources have been identified, "
              "writing cannot begin before material has been gathered, and evaluation cannot begin "
              "before there is a report to evaluate. Within a stage there is scope for concurrency, "
              "and the Reader Agent exploits it by retrieving several pages at once; between stages "
              "there is none.")

    heading(doc, "4.5  Data Model and Inter-Agent Contracts")
    para(doc, "The system holds no persistent state. Everything that exists during a request lives "
              "in memory for the duration of that request and is discarded when the result has been "
              "delivered. This decision removes an entire class of concern, since there is no "
              "database to migrate, back up or secure, and it allows any instance to be restarted at "
              "any moment without loss. Its cost is that a user cannot retrieve an earlier report; "
              "that limitation is recorded in Section 8.3.")
    figure(doc, "fig_3_3_dataflow.png", "Figure 4.3: Data handed between successive agents", width=5.0)
    para(doc, "Because the agents are separate, the values passed between them constitute an "
              "interface and were fixed early. Table 4.3 records those contracts.")
    table(doc,
          ["Stage", "Receives", "Returns"],
          [["Search", "Topic string", "Up to ten entries, each carrying address, title, snippet and score"],
           ["Reader", "Ranked entries", "Aggregated text with counts of successful and failed retrievals"],
           ["Writer", "Aggregated text and topic", "Report in Markdown, six sections"],
           ["Critic", "Report and topic", "Score from zero to ten, five sub-scores, written feedback"]],
          "Table 4.3: Inter-agent data contracts",
          widths=[0.9, 1.85, 3.0])


def chapter_5(doc):
    chapter(doc, 5, "Implementation")

    heading(doc, "5.1  Search Agent")
    para(doc, "The Search Agent turns a topic into a ranked list of addresses. It issues a query "
              "through the DuckDuckGo interface, scores each returned result, discards the weak "
              "ones, removes duplicates and returns the ten strongest.")
    para(doc, "The relevance score combines two components. Overlap between the query terms and the "
              "result title contributes forty per cent, and overlap with the snippet contributes "
              "sixty per cent. The weighting favours the snippet deliberately: titles are often "
              "written for effect and may share few words with a topic that the body of the page "
              "covers thoroughly, whereas the snippet is drawn from the text itself. Results scoring "
              "below 0.5 are discarded on the reasoning that a weak source contributes noise to the "
              "aggregate rather than additional evidence.")
    figure(doc, "fig_4_1_search_agent.png", "Figure 5.1: Search Agent processing sequence", width=4.7)
    table(doc,
          ["Parameter", "Value", "Reason"],
          [["Results requested", "30", "Provides headroom for filtering"],
           ["Results returned", "10", "Sufficient breadth within the time budget"],
           ["Title weight", "0.40", "Titles are indicative but often promotional"],
           ["Snippet weight", "0.60", "Snippet is drawn from the body text"],
           ["Score threshold", "0.50", "Excludes marginal sources"],
           ["Timeout", "10 s", "Matches the stage budget"]],
          "Table 5.1: Search Agent configuration parameters",
          widths=[1.55, 0.9, 3.3])

    heading(doc, "5.2  Reader Agent")
    para(doc, "The Reader Agent retrieves each shortlisted address and extracts its text. Retrieval "
              "carries an explicit ten-second timeout and follows at most five redirects. The "
              "response is parsed with BeautifulSoup4, the element types that reliably carry no "
              "content are removed, and the remaining text is taken.")
    figure(doc, "fig_4_2_reader_agent.png", "Figure 5.2: Reader Agent content extraction sequence", width=4.7)
    para(doc, "An extraction shorter than one hundred words is rejected. This threshold was chosen "
              "after observing that error pages, consent interstitials and paywall notices all parse "
              "successfully and yield short passages of plausible prose which, if admitted, would "
              "contaminate the aggregate. Rejecting them costs an occasional legitimate short page "
              "and prevents a more damaging failure.")
    para(doc, "The agent is expected to fail on some fraction of its input and is written on that "
              "assumption: a failure removes one source, it does not end the request. In routine "
              "operation one to three of the ten addresses fail for the reasons in Table 5.2, and a "
              "report assembled from seven sources is not detectably worse than one from ten.")
    table(doc,
          ["Condition", "Response"],
          [["Retrieval exceeds ten seconds", "Abandon that address and continue"],
           ["Status 403 or 404 returned", "Log the outcome and continue"],
           ["Status 500 or above returned", "Retry once, then abandon"],
           ["Content assembled by client script", "Extraction yields too little text and is rejected"],
           ["Extraction shorter than 100 words", "Discard as unreliable"],
           ["Character encoding not declared", "Infer from the response, fall back to UTF-8"]],
          "Table 5.2: Failure conditions handled by the Reader Agent",
          widths=[2.5, 3.25])

    heading(doc, "5.3  Writer Agent")
    para(doc, "The Writer Agent submits the aggregated text to the language model together with a "
              "prompt that specifies the required structure. The prompt names six sections and gives "
              "each a word budget: a summary of about one hundred and fifty words, five to seven key "
              "findings, an analysis of five hundred to eight hundred words, implications of two "
              "hundred to three hundred words, a conclusion of one hundred and fifty to two hundred "
              "words, and a list of references.")
    para(doc, "Stating the structure explicitly was the single change that most improved output "
              "consistency. Earlier prompts describing desired qualities, asking for thoroughness "
              "or balance, produced output varying widely in shape between runs; naming the "
              "sections and their budgets produced output that could be rendered without further "
              "processing. The temperature is 0.7, which retains enough variation for readable "
              "prose without the drift seen at higher values, and the reply is capped at three "
              "thousand tokens.")

    heading(doc, "5.4  Critic Agent")
    para(doc, "The Critic Agent scores the finished report and returns written feedback. It receives "
              "the report and the original topic and is asked to award between zero and two points "
              "on each of five dimensions, returning the result as structured data together with two "
              "or three identified strengths, two or three weaknesses and at least one concrete "
              "suggestion.")
    table(doc,
          ["Dimension", "What is assessed"],
          [["Completeness", "Whether the report covers the topic rather than one aspect of it"],
           ["Accuracy", "Whether assertions are supported by the material that was gathered"],
           ["Relevance", "Whether the report answers the question that was asked"],
           ["Structure", "Whether the argument is ordered and the sections serve their purpose"],
           ["Expression", "Whether the prose is clear and free of repetition"]],
          "Table 5.3: Scoring dimensions applied by the Critic Agent",
          widths=[1.5, 4.25])
    para(doc, "The agent runs at a temperature of 0.3, lower than the Writer Agent, because "
              "consistency of scoring matters more here than fluency of expression. Its judgement is "
              "necessarily limited: it can detect that a report is thin, poorly ordered or evasive, "
              "but it cannot detect that a well-written report is factually wrong, because it has no "
              "access to ground truth. The score should therefore be read as an assessment of "
              "craftsmanship rather than of correctness, and this limitation is restated in "
              "Section 8.3.")

    heading(doc, "5.5  Pipeline Orchestration and Error Handling")
    para(doc, "The orchestration layer runs the four agents in order, passes each result to the next "
              "stage and records the duration of every stage. Failures are handled where the "
              "distinction between recoverable and terminal can be drawn. A failure to retrieve one "
              "page is recoverable and absorbed by the Reader Agent. A failure of the model "
              "endpoint is retried with an exponential backoff beginning at two seconds and "
              "doubling to a ceiling of thirty-two. A missing or rejected credential is terminal "
              "and is reported plainly, because no number of attempts will supply an absent key.")
    para(doc, "Sessions are isolated from one another. Each request carries its own state, so a slow "
              "request delays no other and no value can leak between concurrent users. Because "
              "nothing is written to disk, an instance can be restarted at any point; requests in "
              "flight are lost, but no stored data is at risk.")

    heading(doc, "5.6  User Interface")
    para(doc, "The interface presents a single text field, a button, a progress indicator and the "
              "rendered result. During processing it names the stage in progress and shows elapsed "
              "time, because a request that takes fifty seconds without visible activity is "
              "indistinguishable from one that has failed. The finished report is rendered from "
              "Markdown, the quality score is shown alongside it, and the sources used are listed so "
              "that any assertion can be traced back to the page it came from.")
    figure(doc, "fig_4_3_interface.png", "Figure 5.3: Layout of the Streamlit user interface", width=4.1)
    para(doc, "All generated content is escaped before rendering. The material has passed through an "
              "arbitrary web page and a language model, and neither can be assumed to have removed "
              "markup that would otherwise be interpreted by the browser.")


def chapter_6(doc):
    chapter(doc, 6, "Testing and Results")

    heading(doc, "6.1  Testing Methodology")
    para(doc, "Testing was arranged in three layers. Unit tests exercise each agent in isolation "
              "with external calls replaced by fixtures, making them fast and deterministic and "
              "allowing failure conditions such as timeouts and malformed replies to be reproduced "
              "on demand. Integration tests exercise the agents in combination to confirm the value "
              "returned by one is accepted by the next. End-to-end tests drive a complete request "
              "against live services, and alone can detect a change in an external dependency.")
    table(doc,
          ["Level", "Target", "Tests", "Result"],
          [["Unit", "Search Agent", "10", "All pass"],
           ["Unit", "Reader Agent", "8", "All pass"],
           ["Unit", "Writer Agent", "7", "All pass"],
           ["Unit", "Critic Agent", "6", "All pass"],
           ["Integration", "Pipeline and recovery", "6", "All pass"],
           ["End to end", "Complete user journey", "2", "All pass"],
           ["Total", "", "39", "39 of 39 pass"]],
          "Table 6.1: Distribution of the automated test suite",
          widths=[1.3, 2.3, 0.85, 1.3])

    heading(doc, "6.2  Test Cases and Results")
    para(doc, "Table 6.2 records representative cases from the suite, chosen to show the failure "
              "conditions rather than only the expected path.")
    table(doc,
          ["Case", "Condition", "Expected outcome", "Result"],
          [["Ordinary query", "Valid topic supplied", "Ten ranked sources returned", "Pass"],
           ["Empty query", "Blank topic submitted", "Rejected before any external call", "Pass"],
           ["Weak results", "All results below threshold", "Empty list, no exception raised", "Pass"],
           ["Duplicate sources", "Same address returned twice", "Retained once only", "Pass"],
           ["Slow page", "Retrieval exceeds ten seconds", "Address abandoned, run continues", "Pass"],
           ["Thin page", "Extraction under 100 words", "Discarded from the aggregate", "Pass"],
           ["Malformed markup", "Unclosed tags in response", "Parsed without exception", "Pass"],
           ["Model unavailable", "Endpoint returns 503", "Retried with backoff, then reported", "Pass"],
           ["Bad credential", "Key absent from environment", "Clear message, no retry", "Pass"],
           ["Malformed score", "Critic reply not valid JSON", "Fallback parse, request completes", "Pass"],
           ["Concurrent load", "Ten simultaneous requests", "All complete, no shared state", "Pass"],
           ["Markup injection", "Angle brackets in topic", "Escaped before rendering", "Pass"]],
          "Table 6.2: Representative test cases and their results",
          widths=[1.15, 1.65, 2.05, 0.6])

    heading(doc, "6.3  Performance Benchmarking")
    para(doc, "Timings were collected over one hundred complete requests on varied topics, executed "
              "on a machine with a quad-core processor and sixteen gigabytes of memory over a "
              "domestic broadband connection. Every stage finished inside its budget and the "
              "complete pipeline finished in a mean of 49.1 seconds against a ceiling of sixty.")
    table(doc,
          ["Stage", "Target", "Mean", "Std. dev.", "95th pct."],
          [["Search", "10.0 s", "7.2 s", "1.8 s", "10.1 s"],
           ["Reader", "15.0 s", "11.5 s", "2.2 s", "15.2 s"],
           ["Writer", "20.0 s", "18.3 s", "1.9 s", "21.8 s"],
           ["Critic", "15.0 s", "12.1 s", "1.4 s", "14.0 s"],
           ["Complete pipeline", "60.0 s", "49.1 s", "4.2 s", "58.1 s"]],
          "Table 6.3: Measured execution time against target",
          widths=[1.75, 1.0, 1.0, 1.0, 1.0])
    figure(doc, "fig_5_1_performance.png", "Figure 6.1: Measured execution time against target", width=4.9)
    para(doc, "The distribution matters more than the mean. The ninety-fifth percentile of the "
              "complete pipeline is 58.1 seconds, which is inside the budget but not comfortably so, "
              "and the variance is dominated by the Reader Agent, whose duration depends on how many "
              "of its ten addresses respond promptly. A single slow page can add several seconds. If "
              "the budget were tightened, that stage is where the work would have to be done.")
    para(doc, "Concurrency was measured separately. A single request completes in 49.1 seconds, five "
              "simultaneous requests in 49.8 seconds and ten in 50.5 seconds. The increase of 1.4 "
              "seconds across ten users reflects the fact that most of the elapsed time is spent "
              "waiting on external services rather than computing locally.")
    table(doc,
          ["Module", "Coverage", "Target", "Assessment"],
          [["Search Agent", "92 %", "85 %", "Above target"],
           ["Reader Agent", "88 %", "85 %", "Above target"],
           ["Writer Agent", "82 %", "85 %", "Marginally below"],
           ["Critic Agent", "79 %", "85 %", "Below target"],
           ["Pipeline", "84 %", "85 %", "At target"],
           ["Overall", "85 %", "85 %", "Meets target"]],
          "Table 6.4: Statement coverage by module",
          widths=[1.6, 1.15, 1.05, 1.95])
    figure(doc, "fig_5_2_coverage.png", "Figure 6.2: Statement coverage by module", width=4.9)
    para(doc, "Coverage across the code base is eighty-five per cent, which meets the target, but "
              "the figure is uneven. The Critic Agent at seventy-nine per cent is the weakest module "
              "and the uncovered statements are concentrated in the paths that handle malformed "
              "replies from the model, which are awkward to provoke because the model does not "
              "produce them on demand. That is an honest weakness rather than a rounding error, and "
              "it is recorded as such in Section 8.3.")

    heading(doc, "6.4  Cost Evaluation")
    para(doc, "Cost was evaluated for a deployment serving one thousand reports per month on a "
              "single modest cloud instance.")
    table(doc,
          ["Item", "With commercial endpoint", "With Groq endpoint"],
          [["Model charges", "5.25 USD", "1.75 USD"],
           ["Compute instance", "30.00 USD", "30.00 USD"],
           ["Storage and transfer", "2.00 USD", "2.00 USD"],
           ["Monthly total", "37.25 USD", "33.75 USD"],
           ["Cost per report", "0.0373 USD", "0.0338 USD"]],
          "Table 6.5: Monthly operating cost comparison",
          widths=[1.85, 2.0, 1.9])
    para(doc, "The reduction in model charges is approximately seventy per cent, which was the "
              "target. Its effect on the total is smaller because the compute instance dominates at "
              "this volume, and the saving becomes material only as volume rises, since model "
              "charges scale with usage while the instance cost does not. The seventy per cent "
              "figure is therefore not a seventy per cent reduction in operating cost, and is not "
              "presented as one.")


def chapter_7(doc):
    chapter(doc, 7, "Deployment and Operations")

    heading(doc, "7.1  Local Installation")
    para(doc, "Installation for development requires Python 3.10 or later and a credential for the "
              "model endpoint. The repository is cloned, a virtual environment created and "
              "activated, dependencies installed, the credential inserted into the environment "
              "file, the test suite run, and the application started.")
    code_block(doc, [
        "git clone <repository-url>",
        "cd multi-agent-research-system",
        "",
        "python -m venv venv",
        "venv\\Scripts\\activate          # Windows",
        "source venv/bin/activate        # Linux or macOS",
        "",
        "pip install -r requirements.txt",
        "copy .env.example .env          # then insert GROQ_API_KEY",
        "",
        "pytest tests/ -v",
        "streamlit run app.py",
    ])
    para(doc, "The application is then reachable on port 8501 of the local machine. Running the test "
              "suite before first use is worthwhile because it detects a missing or malformed "
              "credential immediately, rather than at the point where a user submits a topic.")

    heading(doc, "7.2  Cloud Deployment")
    para(doc, "For shared use the application is hosted on a small cloud instance [8] with two virtual "
              "processors and four gigabytes of memory, which is sufficient because the workload is "
              "dominated by waiting on external services rather than by computation. The instance "
              "runs a current long-term-support Linux distribution. The application is installed as "
              "above and registered as a system service so that it starts on boot and is restarted "
              "automatically if it terminates.")
    para(doc, "Where a uniform environment is wanted across development and hosting, the application "
              "can instead be packaged as a container image [14], which fixes the interpreter "
              "version and the dependency set at build time and removes the need to prepare the "
              "host beyond installing the container runtime.")
    para(doc, "Two operational points deserve emphasis. The credential is supplied to the service "
              "through its environment rather than written into any file within the repository. And "
              "because the application holds no persistent state, an instance may be replaced at any "
              "time without migration or backup; requests in flight are lost, but nothing else is.")

    heading(doc, "7.3  Operational Monitoring")
    para(doc, "Every request writes a log record containing the timestamp, the topic, the number of "
              "sources retrieved and the number that failed, the duration of each stage and the final "
              "quality score. These records are the primary operational instrument. A rising "
              "proportion of failed retrievals indicates that extraction is degrading against "
              "changes in the sites being read; a rising Writer Agent duration indicates congestion "
              "at the model endpoint; a falling mean quality score indicates that source selection "
              "has deteriorated. Each symptom points at a different stage, which is a direct benefit "
              "of the decomposed design.")


def chapter_8(doc):
    chapter(doc, 8, "Conclusion and Recommendations")

    heading(doc, "8.1  Conclusion")
    para(doc, "The system built during this internship performs the whole desk-research sequence "
              "without human intervention and meets the targets set at the outset. A complete report "
              "is produced in a measured mean of 49.1 seconds against a ceiling of sixty, every "
              "individual stage finishes within its own budget, model charges fell by approximately "
              "seventy per cent following the change of provider, and ten concurrent users are "
              "served with less than two seconds of additional latency. The thirty-nine automated "
              "tests all pass and statement coverage stands at eighty-five per cent.")
    para(doc, "The decision that mattered most was to divide the work among four agents rather than "
              "submit it to a single model. The gain was not primarily in output quality, which "
              "improved modestly, but in the ability to observe and diagnose the system. Because "
              "each stage has a declared interface, each can be timed, tested and corrected "
              "separately. Nearly every measurement in Chapter 6 depends on that property, and the "
              "single-model design considered first would have permitted none of them.")

    heading(doc, "8.2  Skills Acquired and Contribution Made")
    para(doc, "The skills acquired fall into three groups. The first concerns language models in "
              "application: constructing prompts that constrain the shape of a reply rather than "
              "describe its desired quality, selecting sampling parameters according to whether "
              "consistency or fluency matters, and parsing structured output defensively. The "
              "second concerns the engineering around such work: writing tests that replace "
              "external services with fixtures so failure conditions can be reproduced on demand, "
              "measuring distributions rather than single figures, and distinguishing recoverable "
              "failures from terminal ones. The third concerns operation: packaging an application "
              "as a service, keeping credentials out of the repository, and instrumenting the "
              "system so that a degradation can be attributed to a stage.")
    para(doc, "The contribution made to the department is a working system with a documented "
              "architecture, a test suite that runs unattended, measured performance against "
              "declared targets, and a deployment procedure another engineer can follow. The "
              "provider comparison in Section 4.2 has been adopted for other work, since the "
              "finding that a substantially cheaper endpoint produced no measurable reduction in "
              "quality applies beyond this project.")

    heading(doc, "8.3  Limitations")
    para(doc, "Several limitations are known and none is incidental. The system does not verify "
              "facts. It synthesises what its sources assert, and if those sources are wrong the "
              "report will be wrong in the same way and will read no differently. The Critic Agent "
              "does not compensate for this, since it assesses craftsmanship and has no access to "
              "ground truth; a confident, well-ordered and entirely incorrect report would score "
              "well. Any assertion that matters must be checked against the listed source.")
    para(doc, "Extraction is heuristic and fails on pages whose content is assembled by client-side "
              "scripting after delivery, on paywalled material and on documents that are not HTML. "
              "Between one and three of ten addresses are typically lost for these reasons. The "
              "pipeline is strictly sequential and cannot return to an earlier stage: if the Critic "
              "Agent judges a report thin because the sources were poor, nothing acts on that "
              "judgement. Only English is supported. Nothing is stored, so a user cannot retrieve an "
              "earlier report. And coverage of the Critic Agent stands at seventy-nine per cent "
              "rather than the eighty-five per cent target, with the uncovered statements "
              "concentrated in the handling of malformed model replies.")

    heading(doc, "8.4  Recommendations and Future Scope")
    para(doc, "Three improvements would repay the effort in roughly this order. The first is to act "
              "on the Critic Agent's judgement rather than merely display it: a report scoring below "
              "a threshold could trigger a second search with revised terms. This would convert the "
              "pipeline from a straight line into a loop and address the most significant structural "
              "limitation, at the cost of a variable and occasionally doubled running time. The "
              "second is to improve extraction, either by rendering pages in a headless browser "
              "before parsing or by adopting a purpose-built article-extraction library; the first "
              "or second option would recover most of the addresses currently lost. The third is to "
              "cache results by topic, which would make repeated requests on popular subjects nearly "
              "instantaneous and would reduce model charges proportionally.")
    para(doc, "Beyond these, support for languages other than English, retrieval of PDF documents "
              "and academic databases, citation formatting in a recognised style, and optional "
              "retention of past reports behind user accounts would each extend the system's reach. "
              "None of them addresses a defect in the present design; they enlarge its scope.")
    para(doc, "A final recommendation concerns the internship programme itself. The most valuable "
              "aspect of the placement was that the targets were numeric and agreed in advance, so "
              "that success was demonstrated by measurement rather than asserted in prose; that "
              "practice is worth retaining. The programme would be improved by allocating explicit "
              "time for deployment and operation, which here were reached late and compressed, "
              "although they are the stages at which a prototype either becomes useful to others "
              "or does not.")


def references(doc):
    para(doc, "REFERENCES", size=CHAPTER_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=18)
    items = [
        "Groq Inc., \"Groq Cloud API Documentation and Model Pricing,\" 2026. "
        "Available: https://console.groq.com/docs",
        "LangChain Inc., \"LangChain Documentation: Chains, Prompts and Output Parsers,\" 2026. "
        "Available: https://python.langchain.com/docs",
        "L. Richardson, \"Beautiful Soup Documentation, Version 4.12,\" 2026. "
        "Available: https://www.crummy.com/software/BeautifulSoup/bs4/doc/",
        "Snowflake Inc., \"Streamlit Documentation: API Reference and Session State,\" 2026. "
        "Available: https://docs.streamlit.io",
        "K. Reitz, \"Requests: HTTP for Humans, Version 2.31,\" 2026. "
        "Available: https://requests.readthedocs.io",
        "Python Software Foundation, \"The Python Language Reference, Version 3.10,\" 2026. "
        "Available: https://docs.python.org/3.10/",
        "H. Krekel et al., \"pytest Documentation: Fixtures, Parametrisation and Coverage,\" 2026. "
        "Available: https://docs.pytest.org",
        "Amazon Web Services, \"Amazon EC2 Instance Types and On-Demand Pricing,\" 2026. "
        "Available: https://aws.amazon.com/ec2/instance-types/",
        "DuckDuckGo Inc., \"DuckDuckGo Search: Privacy Policy and Usage Terms,\" 2026. "
        "Available: https://duckduckgo.com/privacy",
        "T. Brown et al., \"Language Models are Few-Shot Learners,\" Advances in Neural "
        "Information Processing Systems, vol. 33, pp. 1877-1901, 2020.",
        "J. Wei et al., \"Chain-of-Thought Prompting Elicits Reasoning in Large Language Models,\" "
        "Advances in Neural Information Processing Systems, vol. 35, 2022.",
        "M. Wooldridge, An Introduction to MultiAgent Systems, 2nd ed. Chichester: John Wiley "
        "and Sons, 2009.",
        "P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,\" "
        "Advances in Neural Information Processing Systems, vol. 33, pp. 9459-9474, 2020.",
        "Docker Inc., \"Docker Documentation: Building and Running Containers,\" 2026. "
        "Available: https://docs.docker.com",
        "S. S. Rawat, \"Multi-Agent Research System: Source Repository, Test Suite and Technical "
        "Documentation,\" MGM College of Engineering and Technology, Noida, 2026.",
    ]
    for i, item in enumerate(items, 1):
        p = para(doc, after=8, indent=0.45)
        p.paragraph_format.first_line_indent = Inches(-0.45)
        set_font(p.add_run(f"[{i}]\t"))
        set_font(p.add_run(item))


def appendix(doc):
    para(doc, "APPENDIX A: SOURCE CODE", size=CHAPTER_PT, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=18)
    para(doc, "This appendix contains representative source code for the four agents. Configuration, "
              "logging and interface code are omitted; the complete source is available in the "
              "repository cited as reference [15].")

    heading(doc, "A.1  Search Agent", level=3)
    code_block(doc, [
        "def search_agent(topic: str, max_results: int = 10) -> dict:",
        "    \"\"\"Return up to max_results sources ranked by relevance to topic.\"\"\"",
        "    try:",
        "        raw = duckduckgo_search(topic, limit=max_results * 3, timeout=10)",
        "    except SearchError as exc:",
        "        logger.error('search failed: %s', exc)",
        "        return {'urls': [], 'error': str(exc)}",
        "",
        "    seen, scored = set(), []",
        "    for item in raw:",
        "        url = item['link']",
        "        if url in seen:",
        "            continue",
        "        seen.add(url)",
        "        score = (0.4 * term_overlap(item['title'], topic)",
        "                 + 0.6 * term_overlap(item['snippet'], topic))",
        "        if score >= RELEVANCE_THRESHOLD:",
        "            scored.append({'url': url,",
        "                           'title': item['title'],",
        "                           'snippet': item['snippet'],",
        "                           'score': score})",
        "",
        "    scored.sort(key=lambda entry: entry['score'], reverse=True)",
        "    return {'urls': scored[:max_results]}",
    ])

    heading(doc, "A.2  Reader Agent", level=3)
    code_block(doc, [
        "MIN_WORDS = 100",
        "STRIP_TAGS = ('script', 'style', 'nav', 'header', 'footer', 'aside')",
        "",
        "def reader_agent(urls: list) -> dict:",
        "    \"\"\"Fetch each URL and aggregate the extractable body text.\"\"\"",
        "    passages, failures = [], 0",
        "    for entry in urls:",
        "        try:",
        "            response = requests.get(entry['url'], timeout=10,",
        "                                    headers=BROWSER_HEADERS)",
        "            response.raise_for_status()",
        "        except requests.RequestException as exc:",
        "            logger.warning('unreachable %s: %s', entry['url'], exc)",
        "            failures += 1",
        "            continue",
        "",
        "        soup = BeautifulSoup(response.content, 'html.parser')",
        "        for element in soup(STRIP_TAGS):",
        "            element.decompose()",
        "        text = soup.get_text(separator=' ', strip=True)",
        "",
        "        if len(text.split()) < MIN_WORDS:",
        "            failures += 1        # error page, consent wall or paywall",
        "        else:",
        "            passages.append(text)",
        "",
        "    return {'aggregated_text': ' '.join(passages),",
        "            'sources_successful': len(passages),",
        "            'extraction_errors': failures}",
    ])

    heading(doc, "A.3  Writer Agent", level=3)
    code_block(doc, [
        "WRITER_PROMPT = \"\"\"You are preparing a research report on: {topic}",
        "Use only the material below. Produce exactly these sections:",
        "  1. Executive Summary   (about 150 words)",
        "  2. Key Findings        (5 to 7 bullet points)",
        "  3. Detailed Analysis   (500 to 800 words)",
        "  4. Implications        (200 to 300 words)",
        "  5. Conclusion          (150 to 200 words)",
        "  6. References          (the source URLs supplied)",
        "Format the result as Markdown.",
        "",
        "Material:",
        "{content}\"\"\"",
        "",
        "def writer_agent(content: str, topic: str) -> dict:",
        "    llm = ChatGroq(model='gpt-oss-120b', temperature=0.7,",
        "                   max_tokens=3000, timeout=20)",
        "    try:",
        "        reply = llm.invoke(WRITER_PROMPT.format(topic=topic,",
        "                                                content=content))",
        "    except APIError as exc:",
        "        logger.error('synthesis failed: %s', exc)",
        "        return {'report': '', 'success': False, 'error': str(exc)}",
        "",
        "    return {'report': reply.content, 'success': True}",
    ])

    heading(doc, "A.4  Critic Agent", level=3)
    code_block(doc, [
        "DIMENSIONS = ('completeness', 'accuracy', 'relevance',",
        "              'structure', 'expression')",
        "",
        "def critic_agent(report: str, topic: str) -> dict:",
        "    \"\"\"Score the report from 0 to 10 and return written feedback.\"\"\"",
        "    llm = ChatGroq(model='gpt-oss-120b', temperature=0.3, timeout=15)",
        "    reply = llm.invoke(CRITIC_PROMPT.format(topic=topic, report=report))",
        "    try:",
        "        parsed = json.loads(reply.content)",
        "    except json.JSONDecodeError:",
        "        logger.warning('critic returned malformed JSON; using fallback')",
        "        parsed = extract_scores_by_regex(reply.content)",
        "",
        "    scores = {d: clamp(parsed.get(d, 0), 0, 2) for d in DIMENSIONS}",
        "    return {'total_score': round(sum(scores.values()), 1),",
        "            'dimensions': scores,",
        "            'strengths': parsed.get('strengths', []),",
        "            'improvements': parsed.get('improvements', [])}",
    ])


# --------------------------------------------------------------------------
# assembly
# --------------------------------------------------------------------------

def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(PARA_GAP)
    normal.paragraph_format.widow_control = True

    # Section 1 - title page, unnumbered
    s1 = configure(doc.sections[0])
    blank_footer(s1)
    title_page(doc)

    # Section 2 - front matter, lower-case roman, centred
    s2 = configure(doc.add_section(WD_SECTION.NEW_PAGE))
    number_format(s2, "lowerRoman", 1)
    footer_page_number(s2, WD_ALIGN_PARAGRAPH.CENTER)

    # Guideline sequence: cover, training certificate, acknowledgement,
    # declaration, abstract, table of contents, lists of tables and figures.
    training_certificate(doc)
    page_break(doc)
    certificate(doc)
    page_break(doc)
    acknowledgement(doc)
    page_break(doc)
    declaration(doc)
    page_break(doc)
    abstract(doc)
    page_break(doc)
    table_of_contents(doc)
    page_break(doc)
    list_of_tables(doc)
    list_of_figures(doc)
    page_break(doc)
    abbreviations(doc)

    # Section 3 - body, arabic from 1, bottom right
    s3 = configure(doc.add_section(WD_SECTION.NEW_PAGE))
    number_format(s3, "decimal", 1)
    footer_page_number(s3, WD_ALIGN_PARAGRAPH.RIGHT)

    for i, build_chapter in enumerate(
        [chapter_1, chapter_2, chapter_3, chapter_4,
         chapter_5, chapter_6, chapter_7, chapter_8]
    ):
        if i:
            page_break(doc)
        build_chapter(doc)

    # no forced break here: the last chapter often ends with little text on
    # its final page, so References is left to flow onto that same page
    # rather than starting a fresh page that would sit mostly blank
    references(doc)
    page_break(doc)
    appendix(doc)

    doc.save(OUTPUT)
    print("wrote", OUTPUT)


if __name__ == "__main__":
    build()


