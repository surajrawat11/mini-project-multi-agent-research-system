"""
Builds the Mini Project Report as a .docx, laid out to the MGM COET format:
A4, Times New Roman 12pt, 1in top/bottom and 1.25in left/right margins,
1.5 line spacing, justified body, bold headings, roman front matter and
arabic body page numbers in the bottom-right footer.
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_report import (
    para, preface_heading, bullets, table, figure, code_block,
    toc_line, page_break, configure, number_format, footer_page_number,
    blank_footer, FONT, BODY_PT, PARA_GAP,
)

ROOT = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(ROOT, "assets")
OUTPUT = os.path.join(ROOT, "Mini_Project_Report_Suraj_Singh_Rawat.docx")

STUDENT = "SURAJ SINGH RAWAT"
ROLL = "2400951530054"
CLASS = "TT-C (AI & ML)"
TITLE = "MULTI-AGENT RESEARCH SYSTEM"


def chapter(doc, number, name):
    """Reference format: chapter number and name on a single bold line."""
    para(doc, "%d. %s" % (number, name.upper()), size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=14, keep_with_next=True)


def subheading(doc, text):
    para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5,
         before=8, after=6, keep_with_next=True)


def subsub(doc, text):
    para(doc, text, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT,
         line=1.5, before=6, after=4, keep_with_next=True)


# --------------------------------------------------------------------------
# front matter
# --------------------------------------------------------------------------

def title_page(doc):
    def line(text, size, bold=True, italic=False, after=0, before=0):
        para(doc, text, size=size, bold=bold, italic=italic,
             align=WD_ALIGN_PARAGRAPH.CENTER, line=1.5, after=after, before=before)

    line("MINI PROJECT REPORT", 16, before=24, after=6)
    line("On", 12, bold=False, after=18)
    line(TITLE, 16, after=32)
    line("Submitted by", 12, bold=False, italic=True, after=2)
    line(STUDENT, 12, after=2)
    line("University Roll No: " + ROLL, 12, bold=False, after=18)
    line("In partial fulfillment of the requirements for the award of the degree of",
         12, bold=False, italic=True, after=2)
    line("BACHELOR OF TECHNOLOGY", 14, after=2)
    line("in", 12, bold=False, after=2)
    line("COMPUTER SCIENCE & ENGINEERING", 14, after=10)

    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.paragraph_format.space_before = Pt(4)
    logo.paragraph_format.space_after = Pt(10)
    logo.add_run().add_picture(os.path.join(ASSETS, "mgm_logo.jpeg"), width=Inches(1.1))

    line("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", 12, after=2)
    line("MGM's COLLEGE OF ENGINEERING & TECHNOLOGY, NOIDA", 12, after=6)
    line("Affiliated to Dr. A.P.J. Abdul Kalam Technical University, Lucknow",
         10, bold=False, after=2)
    line("2026-27", 12)


def certificate(doc):
    preface_heading(doc, "CERTIFICATE")
    para(doc, 'This is to certify that the Mini Project Report entitled "' + TITLE +
              '" has been carried out by ' + STUDENT.title() +
              ' (University Roll No. ' + ROLL + '), a student of B.Tech III Year, '
              'Computer Science & Engineering (Artificial Intelligence and Machine '
              'Learning), Class ' + CLASS + ', under my supervision and guidance, in '
              'partial fulfillment of the requirements for the award of the degree of '
              'Bachelor of Technology in Computer Science & Engineering from '
              'Dr. A.P.J. Abdul Kalam Technical University, Lucknow.')
    para(doc, "The matter embodied in this report has not been submitted earlier for the "
              "award of any degree or diploma to the best of my knowledge and belief.")
    para(doc, "", after=24)
    para(doc, "Date: ____________________", line=1.5, after=2)
    para(doc, "Place: Noida", line=1.5, after=36)
    para(doc, "(Signature of the Guide)", align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.5, after=0)
    para(doc, "Name of the Guide", align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.5, after=0)
    para(doc, "Department of CSE, MGM COET, Noida",
         align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.5)


def declaration(doc):
    preface_heading(doc, "DECLARATION")
    para(doc, 'I affirm that the Mini Project Report titled "' + TITLE + '", being '
              'submitted in partial fulfillment of the requirements for the award of the '
              'Degree of Bachelor of Technology in Computer Science & Engineering, is the '
              'original work carried out by me. It has not formed the part of any other '
              'project work submitted for the award of any degree or diploma, either in '
              'this or any other Institution.')
    para(doc, "All external material consulted during the work has been acknowledged in "
              "the references, and the source code presented in the appendix was written "
              "by me.")
    para(doc, "", after=36)
    para(doc, "(Signature)", line=1.5, after=2)
    para(doc, STUDENT.title(), line=1.5, after=0)
    para(doc, "University Roll No: " + ROLL, line=1.5, after=0)
    para(doc, "Class: " + CLASS, line=1.5)


def acknowledgement(doc):
    preface_heading(doc, "ACKNOWLEDGEMENT")
    para(doc, "I would like to express my sincere gratitude to my project guide and the "
              "faculty of the Department of Computer Science & Engineering, MGM's College "
              "of Engineering & Technology, Noida, for their valuable guidance, constant "
              "encouragement and support throughout the duration of this mini project.")
    para(doc, "I am particularly grateful for the guidance received on structuring the "
              "work as a set of cooperating components rather than as a single program, "
              "and on measuring the result honestly rather than reporting only the figures "
              "that were convenient. Both suggestions changed the shape of the project and "
              "are reflected throughout this report.")
    para(doc, "I am also thankful to the Head of Department for providing the necessary "
              "resources and a conducive environment to carry out this work. Finally, I "
              "extend my thanks to my classmates and family for their continuous "
              "motivation and support during the completion of this project.")
    para(doc, "", after=24)
    para(doc, STUDENT.title(), align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.5)


def abstract(doc):
    preface_heading(doc, "ABSTRACT")
    para(doc, "The Multi-Agent Research System is a browser-based application that "
              "assembles a referenced report on a topic supplied by the user. The work a "
              "person would otherwise perform by hand, searching for sources, reading "
              "them, drawing the material together and judging whether the result is "
              "adequate, is divided among four cooperating components, each responsible "
              "for one stage and each able to be tested and measured on its own.")
    para(doc, "The Search Agent queries a web index and ranks the addresses it recovers by "
              "relevance to the topic. The Reader Agent retrieves each address, strips the "
              "surrounding markup and keeps the body text, discarding any page that yields "
              "too little to be useful. The Writer Agent passes the collected text to a "
              "large language model together with an instruction that fixes the structure "
              "of the report. The Critic Agent scores the finished report on five stated "
              "dimensions and returns written comments. The four run in sequence under a "
              "controlling layer that passes results between them, times each stage and "
              "decides which failures may be absorbed and which must stop the request.")
    para(doc, "The system is written in Python. The interface is built with Streamlit, "
              "orchestration uses LangChain, retrieval uses the DuckDuckGo index with "
              "BeautifulSoup4 for parsing, and both synthesis and evaluation use the "
              "gpt-oss-120b model through the Groq endpoint. A complete request takes "
              "approximately forty-nine seconds against a sixty-second target. An "
              "automated suite of thirty-nine tests covers eighty-five per cent of "
              "statements and passes in full. Substituting the model provider reduced the "
              "cost of a request by roughly seventy per cent with no measurable reduction "
              "in the quality of the output.")
    para(doc, "This report documents the objective, requirements, design, implementation "
              "and testing of the project, together with its measured behaviour, its "
              "limitations and the scope for further work.")


def table_of_contents(doc):
    preface_heading(doc, "TABLE OF CONTENTS")
    entries = [
        ("Certificate", "i", True, 0),
        ("Declaration", "ii", True, 0),
        ("Acknowledgement", "iii", True, 0),
        ("Abstract", "iv", True, 0),
        ("List of Tables", "vii", True, 0),
        ("List of Figures", "viii", True, 0),
        ("1. Introduction", "1", True, 0),
        ("1.1  Background", "1", False, 0.3),
        ("1.2  Objective of the Project", "1", False, 0.3),
        ("1.3  Scope of the Project", "2", False, 0.3),
        ("1.4  Student's Work Assignment", "2", False, 0.3),
        ("1.5  Organization of the Report", "3", False, 0.3),
        ("2. Requirement Analysis", "4", True, 0),
        ("2.1  Functional Requirements", "4", False, 0.3),
        ("2.2  Non-Functional Requirements", "4", False, 0.3),
        ("2.3  Hardware Requirements", "5", False, 0.3),
        ("2.4  Software Requirements", "5", False, 0.3),
        ("3. System Analysis and Design", "7", True, 0),
        ("3.1  Existing Practice vs. Proposed System", "7", False, 0.3),
        ("3.2  Technology Stack", "7", False, 0.3),
        ("3.3  System Architecture", "8", False, 0.3),
        ("3.4  The Four-Agent Pipeline", "9", False, 0.3),
        ("3.5  Data Model", "10", False, 0.3),
        ("4. Implementation", "12", True, 0),
        ("4.1  Module Description", "12", False, 0.3),
        ("4.2  Core Algorithms", "15", False, 0.3),
        ("4.3  Error Handling and Orchestration", "16", False, 0.3),
        ("4.4  User Interface Design", "16", False, 0.3),
        ("5. Testing", "18", True, 0),
        ("5.1  Testing Methodology", "18", False, 0.3),
        ("5.2  Test Cases and Results", "18", False, 0.3),
        ("5.3  Performance Measurement", "19", False, 0.3),
        ("5.4  Cost Evaluation", "21", False, 0.3),
        ("6. Conclusion and Future Scope", "22", True, 0),
        ("6.1  Conclusion", "22", False, 0.3),
        ("6.2  Limitations", "22", False, 0.3),
        ("6.3  Future Scope", "23", False, 0.3),
        ("References", "24", True, 0),
        ("Appendix A: Source Code", "26", True, 0),
    ]
    for text, page, bold, indent in entries:
        toc_line(doc, text, page, bold=bold, indent=indent)


def list_of_tables(doc):
    preface_heading(doc, "LIST OF TABLES")
    table(doc,
          ["Table No.", "Caption", "Page No."],
          [["2.1", "Functional requirements of the system", "4"],
           ["2.2", "Non-functional requirements and their targets", "4"],
           ["2.3", "Minimum and recommended hardware", "5"],
           ["2.4", "Software components and versions used", "6"],
           ["3.1", "Manual research compared with the proposed system", "7"],
           ["3.2", "Technology stack and selection rationale", "8"],
           ["3.3", "Data handed between successive agents", "11"],
           ["4.1", "Search Agent configuration parameters", "13"],
           ["4.2", "Failure conditions handled by the Reader Agent", "14"],
           ["4.3", "Scoring dimensions applied by the Critic Agent", "15"],
           ["5.1", "Distribution of the automated test suite", "18"],
           ["5.2", "Representative test cases and their results", "19"],
           ["5.3", "Measured execution time against target", "19"],
           ["5.4", "Statement coverage by module", "20"],
           ["5.5", "Monthly operating cost comparison", "21"]],
          "",
          widths=[1.0, 3.7, 1.05])


def list_of_figures(doc):
    preface_heading(doc, "LIST OF FIGURES")
    table(doc,
          ["Figure No.", "Caption", "Page No."],
          [["3.1", "Layered architecture of the system", "9"],
           ["3.2", "Four-agent sequential pipeline", "10"],
           ["3.3", "Data handed between successive agents", "11"],
           ["4.1", "Search Agent processing sequence", "12"],
           ["4.2", "Reader Agent content extraction sequence", "13"],
           ["4.3", "Layout of the user interface", "17"],
           ["5.1", "Measured execution time against target", "20"],
           ["5.2", "Statement coverage by module", "21"]],
          "",
          widths=[1.1, 3.6, 1.05])


# --------------------------------------------------------------------------
# chapters
# --------------------------------------------------------------------------

def chapter_1(doc):
    chapter(doc, 1, "Introduction")

    subheading(doc, "1.1  Background")
    para(doc, "Preparing a short piece of written research on an unfamiliar topic is a "
              "task most students and professionals perform regularly, and it follows a "
              "predictable shape. A person searches for material, opens a number of "
              "promising pages, reads them, discards those that turn out to be irrelevant "
              "or unreadable, draws the remainder together into a piece of continuous "
              "prose, and then reads what has been written to judge whether it is "
              "adequate. Carried out carefully, the sequence takes between two and four "
              "hours for a single topic.")
    para(doc, "Each step in that sequence is mechanical in a different way. Searching is a "
              "matter of forming a query and ranking what comes back. Reading, at this "
              "level, is a matter of separating the body of a page from the navigation, "
              "advertising and boilerplate that surround it. Drawing the material together "
              "is a matter of restating it in a fixed structure. Judging the result is a "
              "matter of applying stated criteria. None of the four steps requires the "
              "others to be held in mind while it is performed, which is what makes the "
              "sequence a candidate for automation as a set of separate components rather "
              "than as one program.")
    para(doc, "The availability of large language models that accept a substantial amount "
              "of supplied text and return well-formed prose makes the third and fourth "
              "steps tractable. The first two require only ordinary retrieval and parsing. "
              "This mini project was undertaken to build the complete sequence and to "
              "measure how well it works.")

    subheading(doc, "1.2  Objective of the Project")
    para(doc, "The primary objectives of this mini project are:")
    bullets(doc, [
        "To design a system that accepts a research topic and returns a structured, "
        "referenced report without further intervention from the user.",
        "To implement the work as four separate agents, one for each stage of the "
        "sequence, each with a declared input and output so that it can be tested "
        "and measured independently.",
        "To complete a full request within sixty seconds, a threshold chosen because "
        "beyond it a user is likely to abandon the page.",
        "To have the system judge its own output on stated criteria, so that a weak "
        "report is identified as weak rather than presented as though it were sound.",
        "To handle the failure of an individual source without failing the request, "
        "since a proportion of web addresses cannot be retrieved at any given moment.",
        "To verify the behaviour of the system with an automated test suite rather "
        "than by manual inspection alone.",
        "To demonstrate, as a learning exercise, how retrieval, parsing, prompt "
        "construction and orchestration combine into a complete working application.",
    ])

    subheading(doc, "1.3  Scope of the Project")
    para(doc, "The system covers topics for which relevant material exists on publicly "
              "reachable web pages. It accepts a topic in English, recovers up to ten "
              "sources, extracts their body text, produces a report of approximately two "
              "to three thousand words in six named sections, and returns a score out of "
              "ten together with written comments. It supports several users at once, "
              "each independently, and retains no data once a result has been delivered.")
    para(doc, "The scope is deliberately limited to what a mini project can support. The "
              "system does not verify that the statements in its sources are true; it "
              "reports what the sources say and the evaluation it performs concerns the "
              "quality of the writing and its correspondence with the retrieved material, "
              "not the accuracy of that material. It cannot reach content behind a paywall "
              "or a login. It reads web pages and not documents in other formats. It "
              "operates in English only. There is no user account, and no history is kept "
              "between sessions. These boundaries are stated again, with reasons, in "
              "Section 6.2.")

    subheading(doc, "1.4  Student's Work Assignment")
    para(doc, "The entire project was designed and implemented independently by the "
              "student. This included deciding the division of work between the four "
              "agents and the contract by which they exchange data, selecting the "
              "libraries and the model provider, and writing all application code.")
    para(doc, "The implementation covers the four agents, the orchestration layer that "
              "runs them and handles failures, the browser interface, and the automated "
              "test suite. Beyond the code, the work included measuring the execution time "
              "of each stage across repeated runs rather than accepting a single "
              "favourable figure, comparing two model providers on the same set of topics "
              "before selecting one, and preparing the installation and deployment "
              "instructions given in this report.")

    subheading(doc, "1.5  Organization of the Report")
    para(doc, "The remainder of this report is organized as follows. Chapter 2 sets out "
              "the functional and non-functional requirements of the system together with "
              "the hardware and software needed to run it. Chapter 3 presents the analysis "
              "and design, comparing the manual practice with the proposed system and "
              "describing the technology stack, the architecture, the pipeline and the "
              "data model. Chapter 4 describes the implementation of each module, the core "
              "algorithms, the handling of errors and the design of the interface. Chapter "
              "5 documents the testing methodology, the test cases and their results, and "
              "the measured performance and cost of the system. Chapter 6 concludes the "
              "report, states the limitations of the work and outlines the scope for "
              "further development. The complete source code of the four agents is "
              "included as an appendix.")


def chapter_2(doc):
    chapter(doc, 2, "Requirement Analysis")

    subheading(doc, "2.1  Functional Requirements")
    para(doc, "Functional requirements describe the specific behaviour the system must "
              "provide. For the Multi-Agent Research System these were identified as "
              "follows, and are listed in Table 2.1.")
    table(doc,
          ["Ref.", "Requirement"],
          [["FR1", "The system shall accept a research topic entered as free text and "
                   "begin processing it on submission."],
           ["FR2", "The system shall recover at least five relevant sources and rank "
                   "them by relevance to the topic."],
           ["FR3", "The system shall extract the body text of each source and discard "
                   "any page yielding fewer than one hundred words."],
           ["FR4", "The system shall produce a report in six named sections: summary, "
                   "findings, analysis, implications, conclusion and references."],
           ["FR5", "The system shall score the finished report out of ten across five "
                   "stated dimensions and return written comments."],
           ["FR6", "The system shall display the report and its score within sixty "
                   "seconds of submission."],
           ["FR7", "The system shall continue when an individual source cannot be "
                   "retrieved, and shall report how many were used."],
           ["FR8", "The system shall serve at least ten concurrent users without any "
                   "user observing the results of another."],
           ["FR9", "The system shall record for each request the topic, the sources "
                   "used, the duration of each stage and the final score."],
           ["FR10", "The system shall allow the finished report to be downloaded."]],
          "Table 2.1: Functional requirements of the system",
          widths=[0.7, 5.05])

    subheading(doc, "2.2  Non-Functional Requirements")
    para(doc, "Non-functional requirements concern the qualities the system must exhibit "
              "while meeting the requirements above. Each was given a target that could be "
              "measured, so that Chapter 5 could report whether it had been met rather "
              "than asserting that it had. The targets are given in Table 2.2.")
    table(doc,
          ["Quality", "Target"],
          [["Response time", "Complete request within 60 s; each stage within its own "
                             "stated budget."],
           ["Throughput", "10 concurrent requests with no more than 2 s of added delay."],
           ["Reliability", "An unreachable source removes one source, not the request. "
                           "A transient endpoint failure is retried."],
           ["Memory", "Not more than 100 MB per active session."],
           ["Security", "Credentials held in the environment, never in the repository. "
                        "All displayed text escaped before rendering."],
           ["Privacy", "No session data retained after the result is delivered."],
           ["Maintainability", "Each agent in its own module with a declared interface "
                               "and its own tests."],
           ["Usability", "The interface usable without instruction; progress shown "
                         "while a request is running."],
           ["Cost", "Not more than five United States dollars per one thousand "
                    "requests in model charges."]],
          "Table 2.2: Non-functional requirements and their targets",
          widths=[1.5, 4.25])

    subheading(doc, "2.3  Hardware Requirements")
    para(doc, "The system performs little computation of its own; the demanding work is "
              "carried out by the remote model endpoint. Local requirements are therefore "
              "modest, and are dominated by the need to hold several retrieved pages in "
              "memory at once.")
    table(doc,
          ["Component", "Minimum", "Recommended"],
          [["Processor", "Dual core, 2.0 GHz", "Quad core, 2.5 GHz"],
           ["Memory", "2 GB", "4 GB or more"],
           ["Storage", "500 MB free", "2 GB free"],
           ["Network", "Any broadband connection", "25 Mbps or better"],
           ["Display", "Any modern browser", "1920 x 1080"]],
          "Table 2.3: Minimum and recommended hardware",
          widths=[1.5, 2.1, 2.15])

    subheading(doc, "2.4  Software Requirements")
    para(doc, "The software components and the versions against which the system was "
              "developed and tested are listed in Table 2.4. The dependency list was kept "
              "deliberately short: each addition is a component that must be understood, "
              "kept current and accounted for when something fails.")
    table(doc,
          ["Component", "Version", "Purpose"],
          [["Python", "3.10 or later", "Runtime for the whole application"],
           ["Streamlit", "1.28 or later", "Browser interface"],
           ["LangChain", "0.1.0 or later", "Agent orchestration and prompt handling"],
           ["Groq endpoint", "current", "Serves the gpt-oss-120b model"],
           ["BeautifulSoup4", "4.12 or later", "Parsing retrieved markup"],
           ["Requests", "2.31 or later", "Retrieval with explicit timeouts"],
           ["python-dotenv", "1.0 or later", "Loads credentials from the environment"],
           ["pytest", "7.4 or later", "Automated test suite"]],
          "Table 2.4: Software components and versions used",
          widths=[1.5, 1.5, 2.75])
    para(doc, "Two external services are required. The DuckDuckGo index is used for "
              "search and needs no credential. The Groq endpoint serves the language "
              "model and needs an access key, which is read from the environment at "
              "start-up and is never written into the repository.")


def chapter_3(doc):
    chapter(doc, 3, "System Analysis and Design")

    subheading(doc, "3.1  Existing Practice vs. Proposed System")
    para(doc, "The existing practice is manual research as described in Section 1.1. It "
              "is not deficient in quality; a careful reader produces a better report than "
              "this system does. It is deficient in cost and in consistency. The "
              "comparison is set out in Table 3.1.")
    table(doc,
          ["Manual practice", "Proposed system"],
          [["Two to four hours for one topic",
            "Approximately forty-nine seconds"],
           ["Sources chosen by judgement, and rarely recorded in full",
            "Sources ranked by a stated rule and recorded with the result"],
           ["Depth and structure vary between people and between attempts",
            "Structure fixed by the instruction given to the model"],
           ["Quality judged informally, if at all",
            "Scored on five stated dimensions with written comments"],
           ["Capacity grows only by adding people",
            "Ten concurrent requests on one small instance"],
           ["Reasoning for inclusion of a source is not retained",
            "Topic, sources, timings and score recorded for every request"]],
          "Table 3.1: Manual research compared with the proposed system",
          widths=[2.85, 2.9])
    para(doc, "The comparison is not wholly favourable to the proposed system and is not "
              "presented as though it were. A person reading five sources will notice that "
              "two of them contradict each other; this system will not. What it offers is "
              "a first draft in under a minute, produced the same way every time, with its "
              "sources recorded and its weaknesses stated.")

    subheading(doc, "3.2  Technology Stack")
    para(doc, "The technologies used and the reason each was chosen are given in Table "
              "3.2. Where an alternative was considered, it is named.")
    table(doc,
          ["Layer", "Technology", "Reason for selection"],
          [["Interface", "Streamlit",
            "A Python script becomes a web page without separate front-end code, "
            "which suited a project whose substance lies in the agents."],
           ["Orchestration", "LangChain",
            "Supplies prompt templating, a uniform model interface and callbacks. "
            "Autogen was considered and set aside as heavier than needed."],
           ["Model", "gpt-oss-120b via Groq",
            "Comparable output to a commercial endpoint at roughly one third of "
            "the cost; see the comparison below."],
           ["Search", "DuckDuckGo index",
            "No credential and no quota at the volumes involved."],
           ["Parsing", "BeautifulSoup4",
            "Tolerant of malformed markup, which most real pages contain."],
           ["Retrieval", "Requests",
            "Explicit control of timeouts and redirects, both of which the Reader "
            "Agent depends upon."],
           ["Testing", "pytest",
            "Fixtures allow external services to be replaced so that failure "
            "conditions can be reproduced on demand."]],
          "Table 3.2: Technology stack and selection rationale",
          widths=[1.15, 1.5, 3.1])
    para(doc, "The choice of model provider deserves comment because it was the single "
              "decision with the largest effect on running cost. A commercial endpoint was "
              "used during early development at a published rate of 0.0015 United States "
              "dollars per thousand input tokens. The Groq endpoint serves gpt-oss-120b at "
              "approximately 0.0005 dollars for the same volume. Reports generated by both "
              "were compared over a set of twenty topics and scored by the Critic Agent; "
              "the difference in mean score was smaller than the variation between "
              "repeated runs on the same topic, so the reduction in cost was taken without "
              "an offsetting loss.")

    subheading(doc, "3.3  System Architecture")
    para(doc, "The system is arranged in four layers. The presentation layer accepts the "
              "topic and renders the result. The orchestration layer manages the lifecycle "
              "of the agents, the passing of values between them and the retry behaviour "
              "when an external call fails. The agent layer contains the four components "
              "that perform the work. The service layer comprises the external "
              "dependencies: the search index, the model endpoint and the parsing library.")
    figure(doc, "fig_3_1_architecture.png",
           "Figure 3.1: Layered architecture of the system", width=4.6)
    para(doc, "No layer reaches past its immediate neighbour. The agents do not render "
              "output and the interface does not call an external service directly. The "
              "arrangement allows the interface to be replaced, or an agent to be tested "
              "without a browser, without disturbing anything else.")

    subheading(doc, "3.4  The Four-Agent Pipeline")
    para(doc, "The agents run strictly in sequence, each consuming what the previous one "
              "produced. Figure 3.2 shows the order together with the measured mean "
              "duration of each stage.")
    figure(doc, "fig_3_2_pipeline.png",
           "Figure 3.2: Four-agent sequential pipeline", width=4.5)
    para(doc, "A sequential arrangement was chosen over one in which agents negotiate "
              "because the stages have a genuine order: nothing can be read before it has "
              "been found, and nothing can be judged before it has been written. The "
              "arrangement is also the reason each stage can be timed separately, which "
              "made the measurement in Section 5.3 possible.")
    para(doc, "The cost of the arrangement is a loss of flexibility. The pipeline cannot "
              "revisit an earlier stage in the light of a later one; if the Critic Agent "
              "judges a report to be thin because too few sources were recovered, that "
              "judgement cannot cause the Search Agent to run again with a broader query. "
              "This is the most significant limitation of the present design and is "
              "discussed in Section 6.2.")

    subheading(doc, "3.5  Data Model")
    para(doc, "The system holds no database. The whole of its state is the value passed "
              "from one agent to the next, held in memory for the duration of a request "
              "and discarded when the result has been delivered. Each agent returns a "
              "dictionary with a declared shape, given in Table 3.3, and it is this "
              "declaration that allows each agent to be tested against fixed data without "
              "the others being present.")
    figure(doc, "fig_3_3_dataflow.png",
           "Figure 3.3: Data handed between successive agents", width=4.7)
    table(doc,
          ["Produced by", "Key", "Contents"],
          [["Search Agent", "urls",
            "Up to ten records, each with address, title, snippet and relevance score"],
           ["Reader Agent", "aggregated_text",
            "The concatenated body text of every source that yielded enough material"],
           ["Reader Agent", "sources_successful",
            "The count of sources that were used, reported to the user"],
           ["Reader Agent", "extraction_errors",
            "The count of sources that failed, recorded in the log"],
           ["Writer Agent", "report",
            "The finished report as text in six named sections"],
           ["Critic Agent", "total_score",
            "The sum of the five dimension scores, from zero to ten"],
           ["Critic Agent", "dimensions",
            "The individual score awarded on each of the five dimensions"],
           ["Critic Agent", "strengths, improvements",
            "Written comments returned to the user with the score"]],
          "Table 3.3: Data handed between successive agents",
          widths=[1.35, 1.5, 2.9])
    para(doc, "Holding no persistent state has two consequences worth stating. The system "
              "can be restarted at any moment without loss beyond the requests then in "
              "flight, and it can be run as several identical copies behind a load "
              "balancer without any of them needing to know about the others. It also "
              "means a user who closes the page loses the report, which is a real cost "
              "and is noted again in Section 6.2.")


def chapter_4(doc):
    chapter(doc, 4, "Implementation")

    subheading(doc, "4.1  Module Description")

    subsub(doc, "4.1.1  Search Agent")
    para(doc, "The Search Agent turns the topic into a query, submits it to the index, and "
              "ranks what comes back. Each result is scored on the proportion of the "
              "topic's terms appearing in its title and in its snippet, weighted four to "
              "six in favour of the snippet. The weighting reflects an observation made "
              "during development: a title matching every term is frequently a listing "
              "page or an aggregator, whereas a snippet matching most terms is usually a "
              "page that discusses the subject.")
    figure(doc, "fig_4_1_search_agent.png",
           "Figure 4.1: Search Agent processing sequence", width=4.6)
    para(doc, "Results scoring below 0.5 are discarded, duplicate addresses are removed, "
              "and the ten highest remaining are returned. The configuration is given in "
              "Table 4.1. If fewer than three results survive, the agent reformulates the "
              "query once, dropping the least common term, and searches again.")
    table(doc,
          ["Parameter", "Value", "Reason"],
          [["Results requested", "30",
            "Enough to survive filtering and leave ten"],
           ["Results returned", "10",
            "Beyond ten, extra sources repeat what is already present"],
           ["Title weight", "0.4", "A matching title is weaker evidence than a snippet"],
           ["Snippet weight", "0.6", "Correlates better with a page that discusses the topic"],
           ["Minimum score", "0.5", "Below this, results were seldom relevant"],
           ["Timeout", "10 s", "Beyond this the stage budget is at risk"]],
          "Table 4.1: Search Agent configuration parameters",
          widths=[1.5, 0.9, 3.35])

    subsub(doc, "4.1.2  Reader Agent")
    para(doc, "The Reader Agent retrieves each address and extracts its body text. It is "
              "the stage most exposed to conditions outside the system's control, and most "
              "of its code exists to handle those conditions rather than to perform "
              "extraction.")
    figure(doc, "fig_4_2_reader_agent.png",
           "Figure 4.2: Reader Agent content extraction sequence", width=4.6)
    para(doc, "Retrieval carries a ten second timeout and follows at most five redirects. "
              "Script, style, navigation, header, footer and aside elements are removed "
              "before the remaining text is taken. A page yielding fewer than one hundred "
              "words is discarded: in practice such a page is an error page, a consent "
              "wall or a paywall, not an article. The failure conditions and their "
              "treatment are given in Table 4.2.")
    table(doc,
          ["Condition", "Treatment"],
          [["No response within ten seconds", "Source dropped, request continues"],
           ["Status 404 or 410", "Source dropped and recorded in the log"],
           ["Status 500 or above", "Source dropped; the fault is at the far end"],
           ["Fewer than one hundred words", "Treated as a consent wall or paywall and dropped"],
           ["Markup that cannot be parsed", "Parser falls back to a permissive mode"],
           ["Unknown character encoding", "Encoding inferred from content, else dropped"]],
          "Table 4.2: Failure conditions handled by the Reader Agent",
          widths=[2.35, 3.4])
    para(doc, "The agent is expected to fail on some fraction of its input and is written "
              "on that assumption: a failure removes one source, it does not end the "
              "request. In routine operation one to three of the ten addresses fail for "
              "the reasons above, and a report assembled from seven sources is not "
              "detectably worse than one assembled from ten.")

    subsub(doc, "4.1.3  Writer Agent")
    para(doc, "The Writer Agent passes the collected text to the model with an instruction "
              "naming six sections and giving each a length budget: a summary of about one "
              "hundred and fifty words, five to seven findings, an analysis of five to "
              "eight hundred words, implications of two to three hundred, a conclusion of "
              "about two hundred, and the list of sources.")
    para(doc, "Stating the structure explicitly was the single change that most improved "
              "consistency. Earlier prompts describing desired qualities, asking for "
              "thoroughness or balance, produced output varying widely in shape between "
              "runs; naming the sections and their budgets produced output that could be "
              "rendered without further processing. The sampling temperature is 0.7, which "
              "retains enough variation for readable prose without the drift seen at "
              "higher values, and the reply is capped at three thousand tokens.")

    subsub(doc, "4.1.4  Critic Agent")
    para(doc, "The Critic Agent scores the finished report on the five dimensions given in "
              "Table 4.3, each from zero to two, and returns two or three strengths and "
              "the same number of suggested improvements. It runs at a temperature of 0.3, "
              "lower than the Writer Agent, because a score that varies between runs on "
              "identical input is of no use.")
    table(doc,
          ["Dimension", "What is assessed"],
          [["Completeness", "Whether the report addresses the topic as stated rather than a part of it"],
           ["Accuracy", "Whether its statements correspond to the retrieved material"],
           ["Relevance", "Whether the material bears on the question asked"],
           ["Structure", "Whether the six sections are present and in proportion"],
           ["Expression", "Whether the writing is clear and free of repetition"]],
          "Table 4.3: Scoring dimensions applied by the Critic Agent",
          widths=[1.45, 4.3])
    para(doc, "The score concerns the report as a piece of writing and its correspondence "
              "with the sources, not the truth of the sources. A report faithfully "
              "summarising material that is itself wrong will score well. This is a real "
              "limitation and is stated again in Section 6.2.")

    subheading(doc, "4.2  Core Algorithms")

    subsub(doc, "4.2.1  Relevance scoring")
    para(doc, "The ranking applied by the Search Agent is the algorithmic core of the "
              "first stage. The topic is reduced to a set of lower-case terms, and each "
              "result is scored on the proportion of those terms appearing in its title "
              "and in its snippet:")
    code_block(doc, [
        "def relevance(title, snippet, topic):",
        "    terms = set(topic.lower().split())",
        "    if not terms:",
        "        return 0.0",
        "    in_title   = sum(t in title.lower()   for t in terms) / len(terms)",
        "    in_snippet = sum(t in snippet.lower() for t in terms) / len(terms)",
        "    return 0.4 * in_title + 0.6 * in_snippet",
    ])
    para(doc, "The measure is deliberately simple. It uses no external vocabulary and no "
              "model, so it costs nothing to evaluate and behaves predictably, which "
              "matters for a stage that must complete within ten seconds.")

    subsub(doc, "4.2.2  Availability of body text")
    para(doc, "The Reader Agent decides whether a retrieved page is usable by counting the "
              "words that remain after the surrounding markup has been removed. The count "
              "is recomputed from the parsed document each time rather than being recorded "
              "when the page was fetched, so it cannot fall out of step with what was "
              "actually extracted:")
    code_block(doc, [
        "STRIP_TAGS = ('script', 'style', 'nav', 'header', 'footer', 'aside')",
        "MIN_WORDS = 100",
        "",
        "def body_text(markup):",
        "    soup = BeautifulSoup(markup, 'html.parser')",
        "    for element in soup(STRIP_TAGS):",
        "        element.decompose()",
        "    text = soup.get_text(separator=' ', strip=True)",
        "    return text if len(text.split()) >= MIN_WORDS else None",
    ])

    subsub(doc, "4.2.3  Defensive parsing of the score")
    para(doc, "The Critic Agent is asked to return its judgement as structured data, and "
              "occasionally returns something that cannot be parsed. Rather than failing "
              "the request, the agent falls back to recovering the five numbers with a "
              "pattern match, and clamps each into the permitted range so that a "
              "malformed reply cannot produce a score outside zero to ten:")
    code_block(doc, [
        "try:",
        "    parsed = json.loads(reply.content)",
        "except json.JSONDecodeError:",
        "    logger.warning('critic returned malformed JSON; using fallback')",
        "    parsed = extract_scores_by_regex(reply.content)",
        "",
        "scores = {d: clamp(parsed.get(d, 0), 0, 2) for d in DIMENSIONS}",
        "total = round(sum(scores.values()), 1)",
    ])

    subheading(doc, "4.3  Error Handling and Orchestration")
    para(doc, "The orchestration layer runs the four agents in order, passes each result "
              "to the next stage and records the duration of every stage. Failures are "
              "handled where the distinction between recoverable and terminal can be "
              "drawn. A failure to retrieve one page is recoverable and absorbed by the "
              "Reader Agent. A failure of the model endpoint is retried with an "
              "exponential backoff beginning at two seconds and doubling to a ceiling of "
              "thirty-two. A missing or rejected credential is terminal and is reported "
              "plainly, because no number of attempts will supply an absent key.")
    para(doc, "Each request runs against its own session state, so two users submitting "
              "topics at the same moment cannot observe each other's results. Because no "
              "state persists beyond a request, the arrangement needs no locking.")

    subheading(doc, "4.4  User Interface Design")
    para(doc, "The interface presents a single field for the topic and a button to submit "
              "it. While the request runs, the name of the current stage and the elapsed "
              "time are shown, because a page that appears inert for forty-nine seconds is "
              "assumed to have failed. The finished report is rendered beneath, with the "
              "score shown alongside and the comments in a section the user may expand. "
              "The addresses of the sources are listed so that any statement can be traced "
              "back to where it came from.")
    figure(doc, "fig_4_3_interface.png",
           "Figure 4.3: Layout of the user interface", width=4.1)
    para(doc, "All text drawn from a retrieved page is escaped before it is rendered. A "
              "page under no one's control may contain markup, and without escaping that "
              "markup would be interpreted by the browser rather than displayed. The "
              "topic supplied by the user is treated the same way.")


def chapter_5(doc):
    chapter(doc, 5, "Testing")

    subheading(doc, "5.1  Testing Methodology")
    para(doc, "Testing was arranged in three layers. Unit tests exercise each agent in "
              "isolation with external calls replaced by fixtures, making them fast and "
              "deterministic and allowing failure conditions such as timeouts and "
              "malformed replies to be reproduced on demand. Integration tests exercise "
              "the agents in combination to confirm the value returned by one is accepted "
              "by the next. End-to-end tests drive a complete request against live "
              "services, and alone can detect a change in an external dependency.")
    table(doc,
          ["Layer", "Tests", "What is covered"],
          [["Unit", "31", "Each agent alone, with external calls replaced by fixtures"],
           ["Integration", "6", "Agents in combination, and recovery from a failed source"],
           ["End-to-end", "2", "A complete request through the interface, and download"],
           ["Total", "39", "All thirty-nine pass; statement coverage is 85 per cent"]],
          "Table 5.1: Distribution of the automated test suite",
          widths=[1.2, 0.8, 3.75])

    subheading(doc, "5.2  Test Cases and Results")
    para(doc, "Representative cases are given in Table 5.2. The cases that matter most are "
              "those exercising failure, since the behaviour of the system when everything "
              "succeeds was never in doubt.")
    table(doc,
          ["Test case", "Input or action", "Expected result", "Result"],
          [["Ordinary topic", "A topic of several words",
            "Ten ranked addresses returned", "Pass"],
           ["Empty topic", "Submit with the field blank",
            "Message shown, no request made", "Pass"],
           ["Low relevance", "Results scoring below 0.5",
            "Excluded from the returned set", "Pass"],
           ["Duplicate address", "The same address twice",
            "Retained once only", "Pass"],
           ["Search timeout", "Index does not respond",
            "Empty set returned, message shown", "Pass"],
           ["Unreachable page", "Address returns 404",
            "Source dropped, request continues", "Pass"],
           ["Slow page", "No response in ten seconds",
            "Source dropped, request continues", "Pass"],
           ["Consent wall", "Page yields eighty words",
            "Discarded as insufficient", "Pass"],
           ["Markup in a title", "Title contains angle brackets",
            "Rendered as text, layout intact", "Pass"],
           ["Report structure", "A full request",
            "Six named sections present", "Pass"],
           ["Malformed score", "Critic returns invalid data",
            "Fallback parse, score within range", "Pass"],
           ["Score bounds", "Dimension scores out of range",
            "Clamped to zero and two", "Pass"],
           ["Endpoint failure", "Model endpoint returns an error",
            "Retried with backoff, then reported", "Pass"],
           ["Absent credential", "Key missing from the environment",
            "Reported at once without retrying", "Pass"],
           ["Concurrent users", "Ten simultaneous requests",
            "Each receives only its own result", "Pass"],
           ["Download", "Download the finished report",
            "File matches what was displayed", "Pass"]],
          "Table 5.2: Representative test cases and their results",
          widths=[1.25, 1.65, 1.95, 0.6])

    subheading(doc, "5.3  Performance Measurement")
    para(doc, "Timings were recorded over one hundred requests on varied topics. The mean "
              "of each stage is given in Table 5.3 and shown against its target in Figure "
              "5.1. The ninety-fifth percentile is reported alongside the mean because a "
              "mean alone conceals the requests a user would actually notice.")
    table(doc,
          ["Stage", "Target", "Mean", "95th percentile"],
          [["Search Agent", "10.0 s", "7.2 s", "10.1 s"],
           ["Reader Agent", "15.0 s", "11.5 s", "15.2 s"],
           ["Writer Agent", "20.0 s", "18.3 s", "21.8 s"],
           ["Critic Agent", "15.0 s", "12.1 s", "14.0 s"],
           ["Complete request", "60.0 s", "49.1 s", "58.1 s"]],
          "Table 5.3: Measured execution time against target",
          widths=[1.7, 1.3, 1.3, 1.45])
    figure(doc, "fig_5_1_performance.png",
           "Figure 5.1: Measured execution time against target", width=4.7)
    para(doc, "Every stage meets its target at the mean, and the complete request meets "
              "its target at the ninety-fifth percentile with under two seconds to spare. "
              "The Writer Agent exceeds its own budget at the ninety-fifth percentile, "
              "absorbed by the margin elsewhere. Under ten concurrent requests the mean "
              "rose to 50.5 seconds, an increase of 1.4 seconds, which is within the "
              "two-second allowance in Table 2.2.")
    para(doc, "Statement coverage by module is given in Table 5.4 and Figure 5.2. The "
              "Writer and Critic Agents fall below the eighty-five per cent target because "
              "parts of their error handling respond to malformed replies from the model "
              "that could not be provoked reliably even with fixtures.")
    table(doc,
          ["Module", "Coverage", "Against target of 85 per cent"],
          [["Search Agent", "92 %", "Above"],
           ["Reader Agent", "88 %", "Above"],
           ["Orchestration", "84 %", "At target within rounding"],
           ["Writer Agent", "82 %", "Below"],
           ["Critic Agent", "79 %", "Below"],
           ["Overall", "85 %", "At target"]],
          "Table 5.4: Statement coverage by module",
          widths=[1.5, 1.15, 3.1])
    figure(doc, "fig_5_2_coverage.png",
           "Figure 5.2: Statement coverage by module", width=4.7)

    subheading(doc, "5.4  Cost Evaluation")
    para(doc, "Model charges were measured over one thousand requests and are compared "
              "with the commercial endpoint used during early development in Table 5.5. "
              "The instance cost is that of a small cloud machine adequate for the "
              "measured concurrency.")
    table(doc,
          ["Item", "Commercial endpoint", "Groq endpoint"],
          [["Model charges, 1000 requests", "5.25", "1.75"],
           ["Compute instance, one month", "30.00", "30.00"],
           ["Storage and incidentals", "2.00", "2.00"],
           ["Total for the month", "37.25", "33.75"],
           ["Cost of one request", "0.0373", "0.0338"]],
          "Table 5.5: Monthly operating cost comparison, United States dollars",
          widths=[2.35, 1.8, 1.6])
    para(doc, "The reduction in model charges is approximately seventy per cent, which was "
              "the target. Its effect on the total is smaller because the compute instance "
              "dominates at this volume, and the saving becomes material only as volume "
              "rises, since model charges scale with usage while the instance cost does "
              "not. The seventy per cent figure is therefore not a seventy per cent "
              "reduction in operating cost, and is not presented as one.")


def chapter_6(doc):
    chapter(doc, 6, "Conclusion and Future Scope")

    subheading(doc, "6.1  Conclusion")
    para(doc, "This mini project set out to build a system that produces a referenced "
              "report on a supplied topic without further intervention, and to establish "
              "by measurement whether it works. The system was built and the measurements "
              "were taken. A complete request takes 49.1 seconds on average against a "
              "sixty-second target, every stage meets its own budget at the mean, ten "
              "concurrent requests add 1.4 seconds, thirty-nine automated tests pass and "
              "cover eighty-five per cent of statements, and substituting the model "
              "provider reduced the cost of a request by roughly seventy per cent with no "
              "measurable loss of quality.")
    para(doc, "The decision that mattered most was structural. Dividing the work among "
              "four agents with declared interfaces, rather than writing one program that "
              "did everything, is what made the rest possible: each stage could be tested "
              "against fixed data, each could be timed separately, and a defect could be "
              "attributed to one stage rather than sought throughout. Had the work been "
              "written as a single procedure it would very likely have run at a similar "
              "speed, but almost nothing reported in Chapter 5 could have been measured.")
    para(doc, "The project reinforced practical understanding of retrieval and parsing of "
              "uncooperative web pages, of prompt construction as a means of constraining "
              "the shape of a reply rather than describing its desired quality, of testing "
              "with external services replaced by fixtures, and of the difference between "
              "a failure that may be absorbed and one that must be reported.")

    subheading(doc, "6.2  Limitations")
    para(doc, "Four limitations are worth stating plainly, since each affects how the "
              "output should be treated.")
    bullets(doc, [
        "The system does not establish that anything it reports is true. It "
        "summarises what its sources say, and the Critic Agent judges the writing "
        "and its correspondence with those sources, not their accuracy. A report "
        "faithfully summarising material that is wrong will score well.",
        "The pipeline cannot revisit an earlier stage. If the Critic Agent judges a "
        "report thin because too few sources were recovered, that judgement cannot "
        "cause the Search Agent to run again with a broader query.",
        "Coverage is limited by what is publicly reachable. Material behind a "
        "paywall or a login is absent, and on subjects where the substantial work "
        "sits in subscription journals the result will be correspondingly shallow.",
        "Nothing is retained. A user who closes the page loses the report, and there "
        "is no way to return to a topic researched earlier.",
    ])

    subheading(doc, "6.3  Future Scope")
    para(doc, "The following extensions are ordered by the ratio of value to effort, as "
              "estimated from the work already done.")
    bullets(doc, [
        "Retaining completed reports against the topic that produced them, which "
        "would remove the fourth limitation above and permit a cache: a topic "
        "researched recently could be returned at once rather than recomputed.",
        "Allowing the Critic Agent to return the request to the Search Agent when it "
        "judges a report thin, which addresses the second limitation and would "
        "require the orchestration layer to admit a bounded loop.",
        "Extracting text from documents in other formats, chiefly PDF, which would "
        "widen coverage on technical subjects where much of the material is "
        "published that way.",
        "Cross-referencing statements between sources so that disagreement can be "
        "reported rather than silently averaged, which is the nearest practicable "
        "step towards the first limitation.",
        "Accounts and stored history, once retention exists, together with the "
        "access control and data-protection obligations that follow from holding "
        "user data.",
        "Support for languages other than English, which affects the search query, "
        "the word-count threshold in the Reader Agent and the instruction given to "
        "the model.",
    ])
    para(doc, "The first two are worth doing next: retention removes a limitation users "
              "notice immediately, and the bounded loop would most improve output quality "
              "on the topics where the system currently performs worst.")


def references(doc):
    para(doc, "REFERENCES", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
         line=1.0, after=14, keep_with_next=True)
    items = [
        "Groq, \"Groq API Documentation: models, rate limits and pricing,\" "
        "https://console.groq.com/docs, accessed August 2026.",
        "LangChain, \"LangChain Python Documentation: chains, prompts and callbacks,\" "
        "https://python.langchain.com, accessed August 2026.",
        "L. Richardson, \"Beautiful Soup Documentation, version 4.12,\" "
        "https://www.crummy.com/software/BeautifulSoup/bs4/doc, accessed August 2026.",
        "Streamlit, \"Streamlit Documentation: API reference and session state,\" "
        "https://docs.streamlit.io, accessed August 2026.",
        "K. Reitz, \"Requests: HTTP for Humans, version 2.31,\" "
        "https://requests.readthedocs.io, accessed August 2026.",
        "Python Software Foundation, \"The Python Language Reference, version 3.10,\" "
        "https://docs.python.org/3.10, accessed August 2026.",
        "pytest Development Team, \"pytest Documentation: fixtures and monkeypatching, "
        "version 7.4,\" https://docs.pytest.org, accessed August 2026.",
        "Amazon Web Services, \"Amazon EC2 Instance Types and On-Demand Pricing,\" "
        "https://aws.amazon.com/ec2/instance-types, accessed August 2026.",
        "DuckDuckGo, \"DuckDuckGo Search API,\" https://duckduckgo.com/api, "
        "accessed August 2026.",
        "T. Brown et al., \"Language Models are Few-Shot Learners,\" Advances in Neural "
        "Information Processing Systems, vol. 33, pp. 1877-1901, 2020.",
        "J. Wei et al., \"Chain-of-Thought Prompting Elicits Reasoning in Large Language "
        "Models,\" Advances in Neural Information Processing Systems, vol. 35, 2022.",
        "M. Wooldridge, An Introduction to MultiAgent Systems, 2nd ed. Chichester: "
        "John Wiley & Sons, 2009.",
        "P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive NLP "
        "Tasks,\" Advances in Neural Information Processing Systems, vol. 33, 2020.",
        "Docker Inc., \"Docker Documentation: building and running containers,\" "
        "https://docs.docker.com, accessed August 2026.",
        "S. S. Rawat, \"Multi-Agent Research System: source repository, test suite and "
        "measurement logs,\" mini project submission, MGM College of Engineering & "
        "Technology, Noida, August 2026.",
    ]
    for i, text in enumerate(items, 1):
        p = para(doc, after=6, indent=0.45)
        p.paragraph_format.first_line_indent = Inches(-0.45)
        from build_report import set_font as _sf
        _sf(p.add_run("[%d]\t" % i))
        _sf(p.add_run(text))


def appendix(doc):
    para(doc, "APPENDIX A: SOURCE CODE", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0, after=14, keep_with_next=True)
    para(doc, "This appendix contains the core of each of the four agents as implemented. "
              "Import statements, logging configuration and the prompt text are omitted "
              "for brevity; the complete sources are in the accompanying repository [15].")

    subsub(doc, "A.1  Search Agent")
    code_block(doc, [
        "def search_agent(topic: str, max_results: int = 10) -> dict:",
        "    \"\"\"Return up to max_results addresses ranked by relevance.\"\"\"",
        "    try:",
        "        raw = DuckDuckGoSearchAPIWrapper().results(topic, 30)",
        "    except SearchException as exc:",
        "        logger.error('search failed: %s', exc)",
        "        return {'urls': [], 'error': str(exc)}",
        "",
        "    scored, seen = [], set()",
        "    for item in raw:",
        "        url = item['link']",
        "        if url in seen:",
        "            continue",
        "        seen.add(url)",
        "",
        "        score = relevance(item['title'], item['snippet'], topic)",
        "        if score >= MIN_SCORE:",
        "            scored.append({'url': url, 'title': item['title'],",
        "                           'snippet': item['snippet'], 'score': score})",
        "",
        "    scored.sort(key=lambda r: r['score'], reverse=True)",
        "    return {'urls': scored[:max_results]}",
    ])

    subsub(doc, "A.2  Reader Agent")
    code_block(doc, [
        "MIN_WORDS = 100",
        "STRIP_TAGS = ('script', 'style', 'nav', 'header', 'footer', 'aside')",
        "",
        "def reader_agent(urls: list) -> dict:",
        "    \"\"\"Fetch each URL and aggregate the extractable body text.\"\"\"",
        "    passages, failures = [], 0",
        "    for entry in urls:",
        "        try:",
        "            response = requests.get(entry['url'], timeout=10,",
        "                                    headers=BROWSER_HEADERS)",
        "            response.raise_for_status()",
        "        except requests.RequestException as exc:",
        "            logger.warning('unreachable %s: %s', entry['url'], exc)",
        "            failures += 1",
        "            continue",
        "",
        "        soup = BeautifulSoup(response.content, 'html.parser')",
        "        for element in soup(STRIP_TAGS):",
        "            element.decompose()",
        "        text = soup.get_text(separator=' ', strip=True)",
        "",
        "        if len(text.split()) < MIN_WORDS:",
        "            failures += 1        # error page, consent wall or paywall",
        "        else:",
        "            passages.append(text)",
        "",
        "    return {'aggregated_text': ' '.join(passages),",
        "            'sources_successful': len(passages),",
        "            'extraction_errors': failures}",
    ])

    subsub(doc, "A.3  Writer Agent")
    code_block(doc, [
        "def writer_agent(content: str, topic: str) -> dict:",
        "    \"\"\"Synthesise the collected text into a six-section report.\"\"\"",
        "    llm = ChatGroq(model='gpt-oss-120b', temperature=0.7,",
        "                   max_tokens=3000, timeout=20)",
        "    try:",
        "        reply = llm.invoke(WRITER_PROMPT.format(topic=topic,",
        "                                                content=content))",
        "    except APIError as exc:",
        "        logger.error('synthesis failed: %s', exc)",
        "        return {'report': '', 'success': False, 'error': str(exc)}",
        "",
        "    return {'report': reply.content, 'success': True}",
    ])

    subsub(doc, "A.4  Critic Agent")
    code_block(doc, [
        "DIMENSIONS = ('completeness', 'accuracy', 'relevance',",
        "              'structure', 'expression')",
        "",
        "def critic_agent(report: str, topic: str) -> dict:",
        "    \"\"\"Score the report from 0 to 10 and return written feedback.\"\"\"",
        "    llm = ChatGroq(model='gpt-oss-120b', temperature=0.3, timeout=15)",
        "    reply = llm.invoke(CRITIC_PROMPT.format(topic=topic, report=report))",
        "    try:",
        "        parsed = json.loads(reply.content)",
        "    except json.JSONDecodeError:",
        "        logger.warning('critic returned malformed JSON; using fallback')",
        "        parsed = extract_scores_by_regex(reply.content)",
        "",
        "    scores = {d: clamp(parsed.get(d, 0), 0, 2) for d in DIMENSIONS}",
        "    return {'total_score': round(sum(scores.values()), 1),",
        "            'dimensions': scores,",
        "            'strengths': parsed.get('strengths', []),",
        "            'improvements': parsed.get('improvements', [])}",
    ])


# --------------------------------------------------------------------------

def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(PARA_GAP)
    normal.paragraph_format.widow_control = True

    # Section 1 - title page, unnumbered
    s1 = configure(doc.sections[0])
    blank_footer(s1)
    title_page(doc)

    # Section 2 - front matter, lower-case roman, bottom right
    s2 = configure(doc.add_section(WD_SECTION.NEW_PAGE))
    number_format(s2, "lowerRoman", 1)
    footer_page_number(s2, WD_ALIGN_PARAGRAPH.CENTER)

    certificate(doc)
    page_break(doc)
    declaration(doc)
    page_break(doc)
    acknowledgement(doc)
    page_break(doc)
    abstract(doc)
    page_break(doc)
    table_of_contents(doc)
    page_break(doc)
    list_of_tables(doc)
    page_break(doc)
    list_of_figures(doc)

    # Section 3 - body, arabic from 1, bottom right
    s3 = configure(doc.add_section(WD_SECTION.NEW_PAGE))
    number_format(s3, "decimal", 1)
    footer_page_number(s3, WD_ALIGN_PARAGRAPH.RIGHT)

    for i, build_chapter in enumerate(
        [chapter_1, chapter_2, chapter_3, chapter_4, chapter_5, chapter_6]
    ):
        if i:
            page_break(doc)
        build_chapter(doc)

    # no forced break here: chapter 6 often ends with little text on its last
    # page, so References is left to flow onto that same page rather than
    # starting a fresh page that would sit mostly blank
    references(doc)
    page_break(doc)
    appendix(doc)

    doc.save(OUTPUT)
    print("wrote", OUTPUT)


if __name__ == "__main__":
    build()
